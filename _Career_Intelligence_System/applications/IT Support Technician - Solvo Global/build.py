#!/usr/bin/env python3
"""Tailored CV + Cover Letter (DOCX + PDF) for IT Support Technician (Solvo Global).
Reuses resumes/build_resumes.py (DOCX) and resumes/pdf_builder.py (PDF). Facts per 00_Master_Resume.md.
Run: python3 build.py
"""
import os, sys
from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
RESUMES = os.path.abspath(os.path.join(HERE, "..", "..", "resumes"))
sys.path.insert(0, RESUMES)
import build_resumes as br
import pdf_builder as pb

cv = {
"headline": "IT Support Technician  |  Windows & Microsoft 365  |  Networking & ITSM  |  Asset Management",
"summary": ("IT support professional with around three years delivering reliable first- and second-line support "
    "for 120+ users across multiple sites at roughly 98% uptime. Experienced preparing and configuring Windows "
    "workstations (OS, drivers, network settings), provisioning Microsoft 365 and Active Directory accounts, "
    "troubleshooting hardware and basic network issues using ping, tracert, and DNS tools, and maintaining IT "
    "asset inventory. Disciplined about ITSM ticket documentation, IT procedures, and clear, customer-focused "
    "communication with technical and non-technical users."),
"skills": [
    ("Windows & Endpoints", "Windows 10/11 setup and imaging, drivers, printers, VoIP, desktop deployment, hardware troubleshooting, onboarding/offboarding"),
    ("Microsoft 365", "Exchange Online, Outlook, Teams, SharePoint, user account creation, Entra ID / Active Directory, MFA, group policy basics"),
    ("Networking", "LAN/WAN, TCP/IP, DNS/DHCP, VPN, Wi-Fi, ping/tracert diagnostics, VLANs, switches, firewalls (Cisco Meraki, MikroTik)"),
    ("ITSM & Service Desk", "1st/2nd-line support, Freshdesk, Zoho Desk, Jira Service Management, ITIL, SLA, incident logging, documentation, ticket closure"),
    ("Asset & Deployment", "IT asset inventory and assignment records, equipment deployment and relocation, structured cabling, multi-site rollout support"),
    ("Security & Certs", "Endpoint protection, patching, MFA; CCNA, Microsoft Azure Fundamentals (AZ-900), Google IT Support"),
],
"experience": [
    {"org":"PATH","title":"Health Informatics / HMIS Officer","dates":"Mar 2025 to Present","loc":"Kisumu/Nyamira, Kenya (Hybrid)","bullets":[
        "Provide IT support alongside HMIS duties for 120+ users, sustaining over 98% uptime across multiple sites.",
        "Administer Microsoft 365 and Active Directory accounts, MFA, and endpoint protection; support Windows workstations and backups.",
        "Troubleshoot LAN/WAN, Wi-Fi, DNS, and VPN connectivity; maintain IT asset records and operational documentation.",
        "Log, track, and close incidents in Freshdesk and Zoho Desk under ITIL/SLA procedures.",
    ]},
    {"org":"PATH","title":"ICT Associate","dates":"Jan 2024 to Mar 2025","loc":"Kisumu, Kenya","bullets":[
        "Delivered 1st- and 2nd-line IT support to 120+ users, achieving around 98% uptime and about 25% faster ticket response under SLA.",
        "Prepared and configured Windows 10/11 workstations, drivers, printers, VoIP, and Microsoft 365 (Outlook, Teams, SharePoint).",
        "Created and managed Active Directory and Microsoft 365 user accounts with group policy and least-privilege access.",
        "Diagnosed network issues (LAN/WAN, DNS, DHCP, VPN, Wi-Fi) on Cisco Meraki and MikroTik infrastructure.",
        "Maintained IT asset inventory and assignment records; supported equipment deployment across sites.",
        "Logged, documented, and closed tickets in Freshdesk and Zoho Desk following ITIL workflows.",
    ]},
    {"org":"PATH","title":"ICT Intern","dates":"Apr 2023 to Dec 2023","loc":"Homa Bay, Kenya","bullets":[
        "Supported LAN/WAN, routers, switches, structured cabling, printers, and VoIP; assisted desktop imaging (PXE boot).",
        "Identified and escalated basic hardware faults; logged incidents in Freshdesk under ITIL and SLA standards.",
        "Contributed to IT documentation and asset tracking during facility deployments.",
    ]},
],
"certs": "CCNA  |  Microsoft Azure Fundamentals (AZ-900)  |  Google IT Support  |  IBM Technical Support  |  GitHub Professional  |  Microsoft Data Analysis",
"achievements": [
    "Supported 120+ users at over 98% uptime with disciplined ITSM ticket handling and documentation.",
    "Provisioned and maintained Microsoft 365 and Active Directory accounts across multi-site operations.",
    "Maintained IT asset inventory and supported equipment deployment and relocation.",
    "Competed in cybersecurity hackathons (Cyberise / Communications Authority of Kenya, NRF).",
],
}

COVER_BODY = [
    "Dear Ms. Bungei,",
    "",
    "I am writing to apply for the IT Support Technician role at Solvo Global. Over the past three years I have provided hands-on IT support for more than 120 users across multiple sites, preparing and configuring Windows workstations, managing Microsoft 365 and Active Directory accounts, troubleshooting hardware and network issues, maintaining IT asset records, and closing tickets through ITSM workflows while keeping systems at over 98% uptime.",
    "Day to day I install and configure Windows OS, drivers, and network settings; provision Microsoft 365 accounts; diagnose connectivity problems using ping, tracert, and DNS tools; and document every incident from logging through resolution. I am disciplined about following IT procedures, keeping asset inventory accurate, and communicating clearly with users who need patient, practical help.",
    "I am currently based in Kisumu and open to relocating to Nairobi for the right on-site opportunity. I am also available for the Monday through Friday 4 PM to 1 AM schedule. I hold a BSc in Applied Computing (Cybersecurity & Digital Forensics), plus the CCNA and Microsoft Azure Fundamentals (AZ-900) certifications.",
    "I would welcome the chance to contribute to Solvo Global's IT support team and grow within your enterprise environment. Thank you for your time and consideration.",
    "",
    "Sincerely,",
    "Meshack Ariri",
]

def build_cover_letter(path):
    doc = Document()
    br.set_base_style(doc)
    br.add_name(doc, "MESHACK ARIRI")
    br.add_title(doc, "IT Support Technician")
    br.add_contact(doc, br.CONTACT)
    for para in COVER_BODY:
        if para == "":
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(para)
        r.font.size = Pt(10.5)
    doc.save(path)
    print("Saved:", path)

if __name__ == "__main__":
    br.build(cv, os.path.join(HERE, "Meshack_Ariri_CV_IT_Support_Technician.docx"))
    build_cover_letter(os.path.join(HERE, "Meshack_Ariri_Cover_Letter_IT_Support_Technician.docx"))
    pb.build_cv_pdf(cv, os.path.join(HERE, "Meshack_Ariri_CV_IT_Support_Technician.pdf"), br.CONTACT)
    pb.build_cover_pdf(COVER_BODY, os.path.join(HERE, "Meshack_Ariri_Cover_Letter_IT_Support_Technician.pdf"), br.CONTACT, "IT Support Technician")
    print("Solvo Global application docs generated (DOCX + PDF).")
