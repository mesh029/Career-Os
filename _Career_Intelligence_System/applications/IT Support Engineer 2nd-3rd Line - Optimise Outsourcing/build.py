#!/usr/bin/env python3
"""Tailored CV + Cover Letter (DOCX + PDF) for 2nd/3rd Line IT Support Engineer (Optimise Outsourcing).
Reuses resumes/build_resumes.py (DOCX) and resumes/pdf_builder.py (PDF). Facts per 00_Master_Resume.md.
Run: python3 build.py
"""
import os, sys
from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
RESUMES = os.path.abspath(os.path.join(HERE, "..", "..", "resumes"))
sys.path.insert(0, RESUMES)
import build_resumes as br
import pdf_builder as pb

cv = {
"headline": "IT Support Engineer (2nd / 3rd Line)  |  Microsoft 365, Entra ID & Intune  |  Networking & Infrastructure  |  Service Desk",
"summary": ("IT support engineer with around three years supporting and administering IT for 120+ users across "
    "multiple sites, handling escalated 2nd and 3rd-line issues through to resolution at roughly 98% uptime. Strong in "
    "Microsoft 365 administration (Exchange Online, SharePoint, Teams, Entra ID), endpoint management with Intune, "
    "networking (VLANs, Wi-Fi, switches, firewalls), backups, and endpoint security. Disciplined about ticket quality, "
    "documentation, and clear communication with both technical and non-technical users, and comfortable owning "
    "problems independently in a fast-paced, multi-client style environment."),
"skills": [
    ("Microsoft 365 & Cloud", "Exchange Online, SharePoint, Teams, Entra ID / Azure AD, Intune, compliance and DLP basics, tenant and email migration support, Microsoft Azure (AZ-900)"),
    ("Endpoint & Devices", "Intune / MDM, Windows 10/11, macOS support, device deployment and imaging, endpoint protection, onboarding/offboarding"),
    ("Networking & Infrastructure", "LAN/WAN, VLANs, Wi-Fi, switches, firewalls (FortiGate, Cisco Meraki MX, MikroTik), VPN, DNS/DHCP, backups and disaster recovery"),
    ("Service Desk & Escalation", "2nd/3rd-line troubleshooting, incident and root-cause analysis, ITIL, SLA, ticketing (Freshdesk, Jira, Zoho Desk), ownership through resolution"),
    ("Documentation & Projects", "SOPs and knowledge base, structured documentation standards (IT Glue / PSA familiarity), small infrastructure and improvement projects"),
    ("Security", "MFA and conditional access, endpoint protection, patch management, data protection; cybersecurity and digital forensics foundation"),
],
"experience": [
    {"org":"PATH","title":"IT Support / Systems Officer","dates":"Mar 2025 to Present","loc":"Kenya (Hybrid)","bullets":[
        "Own escalated 2nd and 3rd-line issues through to resolution for 120+ users, sustaining over 98% uptime with backups, patching, and recovery plans.",
        "Administer Microsoft 365 (Exchange Online, SharePoint, Teams) and Entra ID and Active Directory with MFA and conditional access.",
        "Manage endpoints through Intune-style device management, deployment, and endpoint protection across sites.",
        "Troubleshoot networking (VLANs, Wi-Fi, switches, firewalls: FortiGate, Cisco Meraki MX, MikroTik) and support backups and security platforms.",
        "Maintain high documentation standards (SOPs, knowledge base) and communicate clearly with technical and non-technical stakeholders.",
    ]},
    {"org":"PATH","title":"ICT Associate","dates":"Jan 2024 to Mar 2025","loc":"Kisumu, Kenya","bullets":[
        "Delivered 1st, 2nd and 3rd-line support to 120+ users across sites, reaching around 98% uptime and about 25% faster response under SLA.",
        "Administered Active Directory and Microsoft 365; enforced group policy, MFA, and endpoint protection.",
        "Configured and monitored Cisco Meraki and MikroTik networks; troubleshot LAN/WAN, VPN, Wi-Fi, DNS, and DHCP.",
        "Logged and tracked incidents in Freshdesk and Zoho Desk under ITIL/SLA, escalating appropriately and maintaining ticket quality.",
        "Supported Windows 10/11, printers, VoIP, and conferencing; kept IT asset inventory and knowledge base current.",
    ]},
    {"org":"PATH","title":"ICT Intern","dates":"Apr 2023 to Dec 2023","loc":"Homa Bay, Kenya","bullets":[
        "Supported LAN/WAN, routers, switches, structured cabling, printers, and VoIP; assisted server-room diagnostics.",
        "Logged and tracked incidents in Freshdesk under ITIL and SLA workflows and contributed to IT documentation.",
    ]},
],
"certs": "Microsoft Azure Fundamentals (AZ-900)  |  CCNA  |  Google IT Support  |  IBM Technical Support  |  GitHub Professional  |  Microsoft Data Analysis",
"achievements": [
    "Resolved escalated 2nd/3rd-line issues while keeping systems above 98% uptime across multiple sites.",
    "Administered Microsoft 365 and Entra ID with MFA and conditional access for 120+ users.",
    "Maintained high documentation and ticket-quality standards under ITIL/SLA.",
    "Competed in cybersecurity hackathons (Cyberise / Communications Authority of Kenya, NRF).",
],
}

COVER_BODY = [
    "Dear Hiring Manager,",
    "",
    "I am writing to apply for the 2nd to 3rd Line IT Support Engineer role. Over the past three years I have supported and administered IT for more than 120 users across multiple sites, owning escalated issues through to resolution while keeping systems at over 98% uptime. The technical scope of this role fits my experience well, from Microsoft 365 and Entra ID administration to endpoint management, networking, and infrastructure support.",
    "Day to day I administer Microsoft 365 (Exchange Online, SharePoint, Teams) and Entra ID and Active Directory with MFA and conditional access, manage endpoints and deployments, and troubleshoot networking such as VLANs, Wi-Fi, switches, and firewalls (FortiGate, Cisco Meraki, MikroTik). I support backups and endpoint security, and I hold myself to high ticket-quality and documentation standards, communicating clearly with both technical and non-technical users and escalating appropriately when needed.",
    "I want to be straightforward: my experience has been in-house across a multi-site organisation rather than in a formal MSP, and my strongest platform is Windows, with working macOS support. That said, I am used to juggling multiple environments and stakeholders, I take real ownership of problems, and I ramp quickly on new tools such as Jamf, IT Glue, and PSA systems. My degree specialized in cybersecurity and digital forensics, and I hold the CCNA and Microsoft Azure Fundamentals (AZ-900) certifications.",
    "I would welcome the chance to show how I can contribute to your service desk and infrastructure projects. Thank you for your time and consideration.",
    "",
    "Sincerely,",
    "Meshack Ariri",
]

def build_cover_letter(path):
    doc = Document()
    br.set_base_style(doc)
    br.add_name(doc, "MESHACK ARIRI")
    br.add_title(doc, "IT Support Engineer (2nd / 3rd Line)")
    br.add_contact(doc, br.CONTACT)
    for para in COVER_BODY:
        if para == "":
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(para)
        r.font.size = Pt(10.5)
    doc.save(path)
    print("Saved:", path)

if __name__ == "__main__":
    br.build(cv, os.path.join(HERE, "Meshack_Ariri_CV_IT_Support_Engineer.docx"))
    build_cover_letter(os.path.join(HERE, "Meshack_Ariri_Cover_Letter_IT_Support_Engineer.docx"))
    pb.build_cv_pdf(cv, os.path.join(HERE, "Meshack_Ariri_CV_IT_Support_Engineer.pdf"), br.CONTACT)
    pb.build_cover_pdf(COVER_BODY, os.path.join(HERE, "Meshack_Ariri_Cover_Letter_IT_Support_Engineer.pdf"), br.CONTACT, "IT Support Engineer (2nd / 3rd Line)")
    print("Optimise Outsourcing application docs generated (DOCX + PDF).")
