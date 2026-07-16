#!/usr/bin/env python3
"""Tailored CV + Cover Letter (DOCX + PDF) for IT & Systems Administrator (Bridge Talent).
Reuses resumes/build_resumes.py (DOCX) and resumes/pdf_builder.py (PDF). Facts per 00_Master_Resume.md.
Run: python3 build.py
"""
import os, sys
from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
RESUMES = os.path.abspath(os.path.join(HERE, "..", "..", "resumes"))
sys.path.insert(0, RESUMES)
import build_resumes as br
import pdf_builder as pb

cv = {
"headline": "IT & Systems Administrator  |  Microsoft 365 & Entra ID  |  Endpoint (Intune), Automation & VoIP  |  Remote-ready",
"summary": ("IT and systems administrator with around three years at PATH owning and supporting internal IT systems "
    "for 120+ users across multiple sites. Strong across Microsoft 365 (accounts, Teams, SharePoint, licences, "
    "security), Entra ID and Active Directory with MFA and conditional access, endpoint management, and network and "
    "internet troubleshooting for remote staff. Automates routine administration with PowerShell, Bash, and Python, "
    "and has integrated systems through REST APIs on a full platform built end to end. Comfortable with VoIP (3CX), "
    "ticketing, staff onboarding, and clear systems documentation in a distributed, international setup. Available to "
    "work Australian business hours."),
"skills": [
    ("Microsoft 365 & Identity", "Microsoft 365 admin (Exchange Online, Teams, SharePoint, licences), Entra ID / Azure AD, Active Directory, conditional access, MFA, Group Policy"),
    ("Endpoint & Device Management", "Intune / MDM, Windows 10/11, macOS support, device enrolment, onboarding and offboarding, IT asset inventory"),
    ("Automation & Integration", "PowerShell, Bash and Python scripting (user provisioning, licence management, reporting); REST API and middleware integration between business systems"),
    ("Networking & Connectivity", "LAN/WAN, VPN, Wi-Fi, DNS/DHCP, router configuration, ISP escalation, remote-staff troubleshooting"),
    ("Telephony & Collaboration", "VoIP (3CX, Teams Phone), call-centre and ticketing setup, Zoho Desk, Microsoft Teams, SaaS/CRM administration"),
    ("Security & Documentation", "Endpoint protection, MFA and conditional access, phishing response, SOPs and clear systems documentation, knowledge base"),
],
"experience": [
    {"org":"PATH","title":"ICT / Systems Officer","dates":"Mar 2025 to Present","loc":"Kenya (Hybrid)","bullets":[
        "Own day to day reliability of Microsoft 365, servers, and core systems for 120+ users, sustaining over 98% uptime with backups, patching, and recovery plans.",
        "Administer Entra ID and Active Directory with MFA, conditional access, and least-privilege, and manage endpoints through Intune-style device management.",
        "Automate user provisioning, licence management, and reporting with PowerShell, Bash, and Python, reducing manual effort.",
        "Designed and built a full platform (Next.js, Prisma, MySQL) with 40+ REST APIs that integrate data across systems without manual duplication.",
        "Onboard new staff, set up devices and access, and document systems and processes so the team is never dependent on one person.",
        "Support VoIP (3CX), ticketing, and conference/AV, and coordinate ISPs and vendors for remote-staff connectivity.",
    ]},
    {"org":"PATH","title":"ICT Associate","dates":"Jan 2024 to Mar 2025","loc":"Kisumu, Kenya","bullets":[
        "Delivered 1st and 2nd-line support to 120+ users across sites, reaching around 98% uptime and about 25% faster response under SLA.",
        "Administered Active Directory and Microsoft 365 (Outlook, Teams, SharePoint); enforced group policy, MFA, and endpoint protection.",
        "Configured and monitored Cisco Meraki and MikroTik networks; troubleshot LAN/WAN, VPN, Wi-Fi, DNS, DHCP, and ISP issues.",
        "Set up and maintained VoIP (3CX) and supported CRM/SaaS tools; kept IT asset inventory and knowledge base current.",
    ]},
    {"org":"PATH","title":"ICT Intern","dates":"Apr 2023 to Dec 2023","loc":"Homa Bay, Kenya","bullets":[
        "Supported LAN/WAN, routers, switches, structured cabling, printers, and VoIP; assisted server-room diagnostics.",
        "Logged and tracked incidents in Freshdesk under ITIL and SLA workflows and contributed to IT documentation.",
    ]},
],
"certs": "Microsoft Azure Fundamentals (AZ-900)  |  CCNA  |  Google IT Support  |  IBM Technical Support  |  GitHub Professional  |  Microsoft Data Analysis",
"achievements": [
    "Kept Microsoft 365 and core systems above 98% uptime across multi-site operations.",
    "Automated provisioning and reporting with PowerShell, Bash, and Python, cutting manual work.",
    "Built and integrated a full platform with 40+ REST APIs (Next.js, Prisma, MySQL).",
    "Onboarded and trained 80+ staff and maintained clear systems documentation.",
],
}

COVER_BODY = [
    "Dear Hiring Manager,",
    "",
    "I am writing to apply for the IT & Systems Administrator role supporting your Kenya and Australia teams. Over the past three years at PATH I have owned and supported internal IT systems for more than 120 users across multiple sites, and the day to day of this role maps closely to what I already do: managing Microsoft 365, administering identity and access, supporting remote staff and their connectivity, and keeping systems documented and reliable.",
    "I administer Microsoft 365 (accounts, email, Teams, SharePoint, licences, and security) and Entra ID and Active Directory with MFA and conditional access, and I manage endpoints and onboarding for staff joining across locations. I automate routine work such as user provisioning, licence management, and reporting with PowerShell, Bash, and Python, and I troubleshoot networks and home internet for remote staff, including router configuration and escalating to ISPs. I also set up and support VoIP and ticketing, which fits your plan to build out a customer support call centre.",
    "One thing I would bring beyond standard administration is genuine systems integration. I designed and built a full platform (Next.js, Prisma, MySQL) with more than 40 REST APIs that move data cleanly between systems, so I am comfortable connecting your CRM and business tools through APIs rather than manual duplication. I have also worked with Zoho tools and CRM setups, and I am happy to deepen my Zoho ecosystem knowledge quickly.",
    "I want to be straightforward about experience: I have around three years rather than five, but I have carried broad ownership in that time, from identity and endpoints to networks, automation, and building software. I am comfortable working Australian business hours, have reliable high-speed internet, and am confident communicating with staff and clients in clear, professional English.",
    "I would welcome the chance to show how I can keep your internal systems reliable, secure, and well documented as the team scales. Thank you for your time and consideration.",
    "",
    "Sincerely,",
    "Meshack Ariri",
]

def build_cover_letter(path):
    doc = Document()
    br.set_base_style(doc)
    br.add_name(doc, "MESHACK ARIRI")
    br.add_title(doc, "IT & Systems Administrator")
    br.add_contact(doc, br.CONTACT)
    for para in COVER_BODY:
        if para == "":
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(para)
        r.font.size = Pt(10.5)
    doc.save(path)
    print("Saved:", path)

if __name__ == "__main__":
    br.build(cv, os.path.join(HERE, "Meshack_Ariri_CV_IT_Systems_Admin.docx"))
    build_cover_letter(os.path.join(HERE, "Meshack_Ariri_Cover_Letter_IT_Systems_Admin.docx"))
    pb.build_cv_pdf(cv, os.path.join(HERE, "Meshack_Ariri_CV_IT_Systems_Admin.pdf"), br.CONTACT)
    pb.build_cover_pdf(COVER_BODY, os.path.join(HERE, "Meshack_Ariri_Cover_Letter_IT_Systems_Admin.pdf"), br.CONTACT, "IT & Systems Administrator")
    print("Bridge Talent application docs generated (DOCX + PDF).")
