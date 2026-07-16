#!/usr/bin/env python3
"""Tailored CV + Cover Letter (DOCX + PDF) for IT Systems Administrator + Web Dev (Human Capital, Thika).
Reuses resumes/build_resumes.py (DOCX) and resumes/pdf_builder.py (PDF). Facts per 00_Master_Resume.md.
Run: python3 build.py
"""
import os, sys
from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
RESUMES = os.path.abspath(os.path.join(HERE, "..", "..", "resumes"))
sys.path.insert(0, RESUMES)
import build_resumes as br
import pdf_builder as pb

cv = {
"headline": "IT Systems Administrator  |  Networks, Servers & Cybersecurity  |  Web Development (Next.js, WordPress) & UI/UX",
"summary": ("IT systems administrator and developer with around three years supporting IT infrastructure for 120+ users "
    "across multiple sites at roughly 98% uptime, combined with hands-on web development. Strong across networks, "
    "Windows and Linux servers, system security, and troubleshooting, and equally comfortable building web "
    "applications and sites with Next.js, React, TypeScript, and WordPress, including UI/UX and responsive design. "
    "Designed and built a full platform (Next.js, Prisma, MySQL) with 40+ REST APIs, and applies cybersecurity "
    "fundamentals throughout. Analytical, discreet with sensitive data, and effective working independently under pressure."),
"skills": [
    ("Systems & Network Administration", "Windows/Linux servers, Active Directory, Microsoft 365, LAN/WAN, VLANs, VPN, DNS/DHCP, Cisco Meraki, MikroTik, system performance monitoring"),
    ("Web Development", "Next.js, React, TypeScript, JavaScript, HTML/CSS, Tailwind CSS; REST APIs; Prisma; MySQL/PostgreSQL; Git/GitHub"),
    ("Web Design & UI/UX", "Responsive design, component libraries (shadcn/ui), dashboards and data visualization (Recharts), clean and usable interfaces"),
    ("CMS & ERP", "WordPress and CMS-based sites; ERP familiarity; content and site maintenance"),
    ("Cybersecurity & Data Protection", "Firewalls (FortiGate, Cisco Meraki MX, MikroTik), MFA, endpoint protection, patch management, data protection; cybersecurity and digital forensics foundation"),
    ("Support & Delivery", "1st/2nd-line support, incident and root-cause analysis, backups and disaster recovery, documentation, independent problem-solving under pressure"),
],
"experience": [
    {"org":"PATH","title":"ICT / Systems Officer","dates":"Mar 2025 to Present","loc":"Kenya (Hybrid)","bullets":[
        "Administer and maintain networks, Windows and Linux servers, and system security for 120+ users, sustaining over 98% uptime.",
        "Monitor system performance, run routine maintenance, and manage backups, patching, and disaster recovery.",
        "Apply cybersecurity fundamentals (firewalls, MFA, endpoint protection) and manage user accounts and IT assets.",
        "Diagnose and resolve complex technical issues independently and document systems and processes.",
    ]},
    {"org":"Independent / Freelance","title":"Web Developer & IT Consultant","dates":"2023 to Present","loc":"Remote (project-based)","bullets":[
        "Designed and built a full platform (Next.js, TypeScript, Prisma, MySQL) with 40+ REST APIs, dashboards, and maps for facility reporting, IT assets, and tickets.",
        "Built and maintained websites and digital workflows using Next.js, React, and WordPress, with responsive UI/UX and on-page SEO.",
        "Set up Microsoft 365 / Google Workspace and CRM tools for small businesses, and delivered user training.",
    ]},
    {"org":"PATH","title":"ICT Associate","dates":"Jan 2024 to Mar 2025","loc":"Kisumu, Kenya","bullets":[
        "Delivered 1st and 2nd-line support to 120+ users; achieved around 98% uptime and about 25% faster response under SLA.",
        "Configured and monitored Cisco Meraki and MikroTik networks; administered Active Directory and Microsoft 365 with MFA.",
        "Maintained servers, backups, and endpoint protection; kept IT asset inventory and documentation current.",
    ]},
    {"org":"PATH","title":"ICT Intern","dates":"Apr 2023 to Dec 2023","loc":"Homa Bay, Kenya","bullets":[
        "Supported LAN/WAN, servers, structured cabling, printers, and VoIP; assisted server-room diagnostics.",
        "Logged and tracked incidents in Freshdesk under ITIL and SLA workflows and contributed to documentation.",
    ]},
],
"certs": "CCNA  |  Microsoft Azure Fundamentals (AZ-900)  |  Google IT Support  |  IBM Technical Support  |  GitHub Professional  |  Microsoft Data Analysis",
"achievements": [
    "Built a full platform (Next.js, Prisma, MySQL) with 40+ REST APIs, dashboards, and maps.",
    "Kept networks and servers above 98% uptime across multi-site operations.",
    "Delivered websites and digital workflows with responsive UI/UX for small businesses.",
    "Competed in cybersecurity hackathons (Cyberise / Communications Authority of Kenya, NRF).",
],
}

COVER_BODY = [
    "Dear Hiring Manager,",
    "",
    "I am writing to apply for the IT Systems Administrator position. This role stood out to me because it combines the two things I do best: administering IT infrastructure and building for the web. Over the past three years I have managed networks, Windows and Linux servers, and system security for more than 120 users at over 98% uptime, and alongside this I design and build web applications and sites.",
    "On the systems side, I administer networks (Cisco Meraki, MikroTik, LAN and WAN, VPN, Wi-Fi), maintain servers and backups, apply cybersecurity fundamentals such as firewalls, MFA, and endpoint protection, and resolve complex issues independently. On the development side, I build with Next.js, React, TypeScript, and WordPress, with responsive UI/UX, and I designed and built a full platform (Next.js, Prisma, MySQL) with more than 40 REST APIs, dashboards, and maps. I am comfortable across databases and have ERP familiarity.",
    "I am analytical, discreet with sensitive information, and used to working independently under pressure. My degree specialized in cybersecurity and digital forensics, and I hold the CCNA and Microsoft Azure Fundamentals (AZ-900) certifications.",
    "I would welcome the chance to discuss how my mix of systems administration and web development can add value to your team. Thank you for your time and consideration.",
    "",
    "Sincerely,",
    "Meshack Ariri",
]

def build_cover_letter(path):
    doc = Document()
    br.set_base_style(doc)
    br.add_name(doc, "MESHACK ARIRI")
    br.add_title(doc, "IT Systems Administrator")
    br.add_contact(doc, br.CONTACT)
    for para in COVER_BODY:
        if para == "":
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(para)
        r.font.size = Pt(10.5)
    doc.save(path)
    print("Saved:", path)

if __name__ == "__main__":
    br.build(cv, os.path.join(HERE, "Meshack_Ariri_CV_IT_SysAdmin_WebDev.docx"))
    build_cover_letter(os.path.join(HERE, "Meshack_Ariri_Cover_Letter_IT_SysAdmin_WebDev.docx"))
    pb.build_cv_pdf(cv, os.path.join(HERE, "Meshack_Ariri_CV_IT_SysAdmin_WebDev.pdf"), br.CONTACT)
    pb.build_cover_pdf(COVER_BODY, os.path.join(HERE, "Meshack_Ariri_Cover_Letter_IT_SysAdmin_WebDev.pdf"), br.CONTACT, "IT Systems Administrator")
    print("Thika (Web Dev) application docs generated (DOCX + PDF).")
