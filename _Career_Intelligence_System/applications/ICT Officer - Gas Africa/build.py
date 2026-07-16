#!/usr/bin/env python3
"""Tailored CV + Cover Letter (DOCX + PDF) for ICT Officer (Gas Africa).
Reuses resumes/build_resumes.py (DOCX) and resumes/pdf_builder.py (PDF). Facts per 00_Master_Resume.md.
Run: python3 build.py
"""
import os, sys
from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
RESUMES = os.path.abspath(os.path.join(HERE, "..", "..", "resumes"))
sys.path.insert(0, RESUMES)
import build_resumes as br
import pdf_builder as pb

cv = {
"headline": "ICT Officer  |  Systems & Network Administration  |  IT Support, Microsoft 365 & Cybersecurity  |  Asset & Vendor Management",
"summary": ("ICT professional with around three years at PATH administering and supporting ICT systems, hardware, and "
    "network infrastructure for 120+ users across multiple sites at roughly 98% uptime. Strong in Microsoft 365 and "
    "Active Directory, networking (Cisco Meraki, MikroTik), and cybersecurity fundamentals (firewalls, MFA, endpoint "
    "protection), with hands-on backups and disaster recovery, user-account and IT-asset management, documentation, "
    "and vendor coordination. Known for fast fault diagnosis, independent troubleshooting, and explaining technical "
    "issues clearly to non-technical staff."),
"skills": [
    ("Systems & Network Administration", "Windows/Linux Server, Active Directory, Microsoft 365, LAN/WAN, VLANs, VPN, DNS/DHCP, Cisco Meraki, MikroTik, Wi-Fi, network performance monitoring"),
    ("IT Support & Service Desk", "1st/2nd-line support for 120+ users, fault diagnosis, incident and root-cause analysis, ITIL, SLA, ticketing (Freshdesk/Zoho Desk)"),
    ("Cybersecurity & Data Protection", "Firewalls (FortiGate, Cisco Meraki MX, MikroTik), endpoint protection, MFA, patch management, data protection; cybersecurity and digital forensics foundation"),
    ("Backup, Continuity & Maintenance", "System backups and disaster recovery (Veeam/rsync), software updates and patching, routine preventive maintenance, server-room management"),
    ("Assets, Documentation & Vendors", "User-account administration, IT asset tracking and inventory, SOPs and system documentation, vendor and service-provider coordination"),
    ("Delivery & Communication", "Staff training and onboarding, clear communication with non-technical users, prioritising in busy environments"),
],
"experience": [
    {"org":"PATH","title":"ICT / Systems Officer","dates":"Mar 2025 to Present","loc":"Kenya (Hybrid)","bullets":[
        "Administer and maintain ICT systems, servers (Windows and Linux), and network infrastructure for continuous uptime above 98%.",
        "Implement data security, system backups, and disaster recovery, and oversee timely patching and software updates.",
        "Manage user accounts and access controls, track IT assets, and keep system documentation and SOPs current.",
        "Monitor network performance and run routine maintenance checks to prevent failures, coordinating vendors and ISPs when needed.",
        "Diagnose and resolve complex technical issues independently and explain solutions clearly to non-technical staff.",
    ]},
    {"org":"PATH","title":"ICT Associate","dates":"Jan 2024 to Mar 2025","loc":"Kisumu, Kenya","bullets":[
        "Delivered 1st and 2nd-line support to 120+ users across sites, reaching around 98% uptime and about 25% faster response under SLA.",
        "Configured and monitored Cisco Meraki and MikroTik networks; troubleshot LAN/WAN, VPN, Wi-Fi, DNS, and DHCP.",
        "Administered Active Directory and Microsoft 365; enforced least-privilege access, group policy, MFA, and endpoint protection.",
        "Coordinated with vendors and service providers; kept the IT asset inventory and knowledge base current.",
        "Supported endpoint protection, patching, and backups, and installed and maintained Windows PCs and VoIP devices.",
    ]},
    {"org":"PATH","title":"ICT Intern","dates":"Apr 2023 to Dec 2023","loc":"Homa Bay, Kenya","bullets":[
        "Supported LAN/WAN infrastructure, routers, switches, structured cabling, printers, and VoIP; assisted server-room diagnostics.",
        "Logged and tracked incidents in Freshdesk under ITIL and SLA workflows and contributed to IT documentation and audits.",
    ]},
],
"certs": "CCNA  |  Microsoft Azure Fundamentals (AZ-900)  |  Google IT Support  |  IBM Technical Support  |  GitHub Professional  |  Microsoft Data Analysis",
"achievements": [
    "Kept mission-critical systems and networks above 98% uptime across multi-site operations.",
    "Strengthened data security and continuity through backups, patching, and disaster recovery.",
    "Maintained accurate IT asset registers and clear system documentation.",
    "Trained 80+ staff and coordinated vendors for reliable, secure IT operations.",
],
}

COVER_BODY = [
    "Dear Hiring Manager,",
    "",
    "I am writing to apply for the ICT Officer position at Gas Africa. Over the past three years at PATH I have administered and supported ICT systems, hardware, and network infrastructure for more than 120 users across multiple sites, keeping mission-critical systems running at over 98% uptime. The responsibilities in your advert match my daily work closely, from systems and network administration to technical support, data security, and asset management.",
    "I administer Microsoft 365 and Active Directory, configure and monitor networks (Cisco Meraki and MikroTik, LAN and WAN, VPN, Wi-Fi), and apply cybersecurity fundamentals such as firewalls, MFA, and endpoint protection. I manage system backups and disaster recovery, oversee patching and software updates, maintain user accounts and IT asset registers, and keep system documentation current. I also monitor network performance and run routine maintenance to prevent failures, and I coordinate with external vendors and service providers when specialist support is needed.",
    "Colleagues rely on me to diagnose and resolve complex issues independently and to explain technical matters clearly to non-technical staff, which keeps day to day operations smooth in busy environments. My degree specialized in cybersecurity and digital forensics, and I hold the CCNA and Microsoft Azure Fundamentals (AZ-900) certifications.",
    "I would welcome the chance to discuss how my experience can help keep Gas Africa's technology reliable, secure, and well supported. Thank you for your time and consideration.",
    "",
    "Sincerely,",
    "Meshack Ariri",
]

def build_cover_letter(path):
    doc = Document()
    br.set_base_style(doc)
    br.add_name(doc, "MESHACK ARIRI")
    br.add_title(doc, "ICT Officer")
    br.add_contact(doc, br.CONTACT)
    for para in COVER_BODY:
        if para == "":
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(para)
        r.font.size = Pt(10.5)
    doc.save(path)
    print("Saved:", path)

if __name__ == "__main__":
    br.build(cv, os.path.join(HERE, "Meshack_Ariri_CV_ICT_Officer_GasAfrica.docx"))
    build_cover_letter(os.path.join(HERE, "Meshack_Ariri_Cover_Letter_ICT_Officer_GasAfrica.docx"))
    pb.build_cv_pdf(cv, os.path.join(HERE, "Meshack_Ariri_CV_ICT_Officer_GasAfrica.pdf"), br.CONTACT)
    pb.build_cover_pdf(COVER_BODY, os.path.join(HERE, "Meshack_Ariri_Cover_Letter_ICT_Officer_GasAfrica.pdf"), br.CONTACT, "ICT Officer")
    print("Gas Africa application docs generated (DOCX + PDF).")
