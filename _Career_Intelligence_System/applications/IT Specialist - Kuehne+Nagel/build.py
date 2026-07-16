#!/usr/bin/env python3
"""Tailored CV + Cover Letter (DOCX + PDF) for IT Specialist, Kuehne+Nagel.
Facts per 00_Master_Resume.md. Human tone, no em-dashes. Run: python3 build.py
"""
import os, sys
from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
RESUMES = os.path.abspath(os.path.join(HERE, "..", "..", "resumes"))
sys.path.insert(0, RESUMES)
import build_resumes as br
import pdf_builder as pb

cv = {
"headline": "IT Specialist  |  Microsoft Environments & Infrastructure  |  Incident Support, Escalation & Process Improvement",
"summary": ("IT professional with about three years supporting technology infrastructure for 120+ users "
    "across multiple sites at PATH. Handles user and system incidents through trouble tickets, setup and "
    "configuration requests, and escalated issues that need clear diagnostics and handover notes. Strong "
    "in Microsoft 365, Entra ID and Active Directory, endpoint protection, backups, and network "
    "troubleshooting. Improved support response times by around 25% through better ITIL workflows, service "
    "metrics, and knowledge-base documentation. Mentors junior staff, contributes to improvement projects, "
    "and works well with technical and non-technical colleagues across locations. BSc Applied Computing; "
    "AZ-900 and CCNA certified."),
"skills": [
    ("Microsoft Environments", "Microsoft 365 (Exchange Online, Teams, SharePoint, OneDrive), Entra ID / Azure AD, Active Directory, Group Policy, MFA, Windows 10/11, Intune-style endpoint management"),
    ("Incident & Ticket Support", "Trouble-ticket handling, 1st/2nd-line escalation, ITIL and SLA, Freshdesk, Zoho Desk, Jira Service Management, root-cause analysis, error diagnostics and handover documentation"),
    ("Security, AV & Backups", "Endpoint protection and anti-virus, patch management, MFA, IAM, Veeam and rsync backups, disaster recovery, security-awareness training"),
    ("Infrastructure & Networking", "LAN/WAN, VLANs, VPN, DNS/DHCP, Wi-Fi, Cisco Meraki, MikroTik, printers, VoIP (3CX), server-room support"),
    ("Process Improvement & Delivery", "Service metrics and SLA reporting, knowledge base and SOPs, user onboarding, mentoring, small technical improvement projects, PowerShell/Bash/Python automation"),
    ("Collaboration", "Multi-site support, stakeholder communication, training and presentations, cross-cultural teamwork in international program settings"),
],
"experience": [
    {"org":"PATH","title":"ICT / Systems Officer","dates":"Mar 2025 to Present","loc":"Kenya (Hybrid)","bullets":[
        "Handle escalated incidents and system issues for 120+ users, keeping core infrastructure above 98% uptime through patching, backups, and recovery planning.",
        "Administer Microsoft 365 and Entra ID / Active Directory (MFA, access rights, Group Policy) and support setup and change requests from users and site teams.",
        "Prepare clear error diagnostics and incident notes for handover when issues need vendor or senior-engineer input.",
        "Maintain Windows and Linux application servers, endpoint protection, and backup routines used by program operations.",
        "Contribute to improvement projects including automated reporting and operational dashboards that cut manual processing time by about 40%.",
    ]},
    {"org":"PATH","title":"ICT Associate","dates":"Jan 2024 to Mar 2025","loc":"Kisumu, Kenya","bullets":[
        "Resolved user and system incidents via Freshdesk and Zoho Desk under ITIL and SLA, often handling 20+ tickets per day at peak periods.",
        "Completed setup and configuration work with users: Microsoft 365 accounts, devices, printers, VoIP extensions, and network access.",
        "Took on escalated tasks from frontline support, troubleshooting identity, endpoint, and connectivity issues through to resolution.",
        "Studied ticket trends and tightened workflows, improving average response time by about 25% and keeping the knowledge base current.",
        "Supported networked anti-virus, patching, and backup processes alongside Cisco Meraki network monitoring.",
        "Mentored ICT interns and documented standard procedures so the team could repeat common fixes without reinventing steps.",
    ]},
    {"org":"PATH","title":"ICT Intern","dates":"Apr 2023 to Dec 2023","loc":"Homa Bay, Kenya","bullets":[
        "Supported LAN/WAN rollout, desktop imaging (PXE), printers, and VoIP, and logged incidents in Freshdesk with SLA tracking.",
        "Assisted with server-room checks and contributed to IT documentation used by the wider support team.",
    ]},
],
"certs": "Microsoft Azure Fundamentals (AZ-900)  |  CCNA  |  Google IT Support  |  IBM Technical Support  |  GitHub Professional  |  Microsoft Data Analysis",
"achievements": [
    "Supported 120+ users across multiple sites at over 98% uptime in Microsoft-centric environments.",
    "Improved average ticket response time by about 25% through process review and better documentation.",
    "Mentored interns and trained 80+ staff on secure system use across 15+ sessions.",
    "Delivered network setup and migration project from planning through to user handover.",
    "Competed in national cybersecurity hackathons (Cyberise / Communications Authority of Kenya, NRF).",
],
}

COVER_BODY = [
    "Dear Hiring Manager,",
    "",
    "I am applying for the IT role at Kuehne+Nagel. Over the past three years at PATH I have supported and administered technology for more than 120 users across several sites, working through trouble tickets, setup and configuration requests, and escalated issues that needed proper diagnostics before handover or resolution. I am comfortable in Microsoft environments (Microsoft 365, Entra ID and Active Directory, Teams, SharePoint), networked anti-virus and endpoint protection, and backup and recovery routines.",
    "What I enjoy most is closing the loop: finding the actual cause, documenting it clearly, and leaving the user with a fix that lasts. I improved average ticket response time by about 25% by tightening ITIL workflows, keeping the knowledge base current, and tracking simple service metrics for the team. I have also mentored interns and trained 80+ colleagues on secure system use, which fits your emphasis on mentoring and better customer service.",
    "I hold a BSc in Applied Computing (Cybersecurity and Digital Forensics) from KCA University and certifications including Microsoft Azure Fundamentals (AZ-900) and CCNA. I have not worked directly in logistics or transport yet, but I am drawn to Kuehne+Nagel because the work sits behind real-world delivery that matters, and my background is already in keeping multi-site, time-sensitive operations online at around 98% uptime.",
    "I am organised under pressure, happy to work flexible hours, and able to travel when the role requires it. I would welcome the opportunity to discuss how I can support your global IT team and contribute to improvement projects alongside colleagues in other regions.",
    "",
    "Thank you for your time and consideration.",
    "",
    "Sincerely,",
    "Meshack Ariri",
]

def build_cover_letter(path):
    doc = Document()
    br.set_base_style(doc)
    br.add_header(doc, "MESHACK ARIRI", "IT Specialist", br.CONTACT)
    for para in COVER_BODY:
        if para == "":
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(para)
        r.font.size = Pt(10.5)
    doc.save(path)
    print("Saved:", path)

if __name__ == "__main__":
    base = "Meshack_Ariri"
    br.build(cv, os.path.join(HERE, f"{base}_CV_IT_Specialist_Kuehne_Nagel.docx"))
    build_cover_letter(os.path.join(HERE, f"{base}_Cover_Letter_IT_Specialist_Kuehne_Nagel.docx"))
    pb.build_cv_pdf(cv, os.path.join(HERE, f"{base}_CV_IT_Specialist_Kuehne_Nagel.pdf"), br.CONTACT)
    pb.build_cover_pdf(
        COVER_BODY,
        os.path.join(HERE, f"{base}_Cover_Letter_IT_Specialist_Kuehne_Nagel.pdf"),
        br.CONTACT,
        "IT Specialist",
    )
    print("Kuehne+Nagel application docs generated (DOCX + PDF).")
