#!/usr/bin/env python3
"""Tailored CV + Cover Letter (DOCX + PDF) for ICT Officer — USIU-Africa.
Facts per resumes/00_Master_Resume.md. Run: python3 build.py
"""
import os
import shutil
import sys

from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
RESUMES = os.path.abspath(os.path.join(HERE, "..", "..", "resumes"))
ROOT = os.path.abspath(os.path.join(HERE, "..", ".."))
sys.path.insert(0, RESUMES)
import build_resumes as br
import pdf_builder as pb

CONTACT_NAIROBI = (
    "Nairobi / Kisumu, Kenya  |  +254 741 174 779  |  aririmeshack@gmail.com  |  "
    "linkedin.com/in/meshack-ariri  |  meshreallycodes.online"
)

cv = {
"headline": "ICT Officer  |  Service Desk · Microsoft 365 · Networking · ITSM  |  University & Multi-Site Support",
"summary": (
    "ICT professional with ~3 years at PATH delivering first- and second-line technical support and "
    "systems administration for 120+ users across multi-site programs. Strong in Windows 10/11, "
    "Microsoft 365 and Active Directory administration, networking (TCP/IP, DNS, DHCP, VLANs), "
    "ITIL-style ticketing (Freshdesk/Zoho/Jira), asset inventory, and security basics (MFA, patching, "
    "endpoint protection). Sustained ~98% uptime under SLA discipline; trained staff and maintained "
    "SOPs/knowledge bases. Seeking an ICT Officer role at USIU-Africa to support service delivery, "
    "infrastructure availability, and excellent user experience for staff and students."
),
"skills": [
    ("Service Desk & ITSM", "1st-line support, walk-in/phone/email/ticket channels, Freshdesk, Zoho Desk, Jira Service Management, ITIL, SLA, incident logging, knowledge base"),
    ("Windows & Microsoft 365", "Windows 10/11, Microsoft 365 Admin (Outlook, Teams, SharePoint), Active Directory / Entra ID, MFA, onboarding/offboarding, Group Policy basics"),
    ("Networking & Infrastructure", "LAN/WAN, TCP/IP, DNS, DHCP, VLANs, VPN, Wi-Fi, Cisco Meraki, switches/routers, monitoring and escalation"),
    ("Security & Compliance", "Endpoint protection, antivirus/patching, MFA/IAM, least-privilege access, security awareness, access logging; Kenya DPA awareness"),
    ("Assets, AV & Continuity", "ICT asset inventory/tagging, license tracking (M365), backups & DR logs, VoIP/AV and training-venue setup support"),
    ("Enterprise apps & data", "Rapid adoption of line-of-business systems; KenyaEMR/DHIS2 support experience; SQL reporting; documentation/SOPs"),
],
"experience": [
    {"org": "PATH", "title": "Health Informatics / HMIS Officer (ICT support retained)", "dates": "Mar 2025 – Present", "loc": "Kisumu/Nyamira, Kenya (Hybrid)", "bullets": [
        "Provide ICT and systems support for 120+ users alongside HMIS duties, sustaining >98% uptime for critical services.",
        "Administer Microsoft 365 and Active Directory accounts, MFA, and least-privilege access; support Windows endpoints and server backups.",
        "Monitor systems and escalate infrastructure or security incidents promptly; maintain operational logs and documentation.",
        "Support user enablement on enterprise digital tools; train staff and update SOPs/knowledge-base articles.",
        "Maintain IT asset awareness and assist with equipment accountability across program sites.",
    ]},
    {"org": "PATH", "title": "ICT Associate", "dates": "Jan 2024 – Mar 2025", "loc": "Kisumu, Kenya", "bullets": [
        "Served as a primary contact for ICT support requests; resolved Level 1/2 incidents (accounts, software, connectivity) under ITIL/SLA workflows.",
        "Delivered support to 120+ users with ~98% uptime and ~25% faster ticket response through disciplined queue management (Freshdesk/Zoho Desk).",
        "Executed user onboarding/offboarding: Active Directory and Microsoft 365 provisioning, access rights, and deprovisioning for audit-ready access control.",
        "Configured and monitored Cisco Meraki network devices; troubleshot LAN/WAN, DNS, DHCP, VPN, and Wi-Fi issues.",
        "Supported endpoint protection, patching, and MFA; performed health checks on PCs, printers, and related peripherals.",
        "Maintained IT asset inventory records; authored knowledge-base articles and delivered user training.",
    ]},
    {"org": "PATH", "title": "ICT Intern", "dates": "Apr 2023 – Dec 2023", "loc": "Homa Bay, Kenya", "bullets": [
        "Supported LAN/WAN, routers, switches, structured cabling, printers, VoIP, and desktop imaging (PXE).",
        "Logged and tracked Freshdesk tickets under SLA monitoring; contributed to process documentation and internal IT audits.",
        "Assisted facility digitalization deployments and day-to-day ICT troubleshooting for non-technical users.",
    ]},
],
"certs": (
    "CCNA  |  Microsoft Azure Fundamentals (AZ-900)  |  Google IT Support  |  "
    "IBM Technical Support  |  GitHub Professional  |  Microsoft Data Analysis"
),
"achievements": [
    "Supported 120+ users at >98% uptime with SLA-driven service desk discipline — aligned to high-availability service goals.",
    "Improved ticket response ~25% and cut reporting/processing effort ~40% through structured workflows and automation.",
    "Trained 80+ staff on digital tools, secure use, and data quality — strong campus user-enablement signal.",
    "Cybersecurity hackathon participant (Cyberise / CA Kenya; NRF), reinforcing security-awareness mindset.",
],
"referees": [
    "Patrick Ondieki — ICT Officer, PATH  |  pondieki@path.org",
    "Marvin Ngosa — Systems Administrator, PATH  |  mngosa@path.org",
    "Gadaffi Ochieng — Procurement Officer, Ampath  |  ochienggadaffi@gmail.com",
],
}

# DOCX / PDF referees (plain text)
cv_docx_referees = [
    "Patrick Ondieki — ICT Officer, PATH  |  pondieki@path.org",
    "Marvin Ngosa — Systems Administrator, PATH  |  mngosa@path.org",
    "Gadaffi Ochieng — Procurement Officer, Ampath  |  ochienggadaffi@gmail.com",
]

COVER_BODY = [
    "Dear Director Human Resource,",
    "",
    "I am writing to apply for the ICT Officer positions at USIU-Africa "
    "(Service Delivery and Multimedia).",
    "",
    "I have about three years of hands-on ICT experience at PATH. I started as an "
    "ICT Intern in Homa Bay, then moved to ICT Associate, and later supported health "
    "information systems while still doing day to day IT support. That path taught me "
    "one thing clearly: good ICT work is about keeping people working, not just fixing machines.",
    "",
    "Day to day I have supported over 120 users. I handle Level 1 and Level 2 issues "
    "like Microsoft 365 and Active Directory accounts, Windows computers, printers, "
    "Wi-Fi, and VPN problems. I use ticketing tools such as Freshdesk and Zoho, follow "
    "SLAs, and we kept systems at about 98% uptime. I also do user onboarding and "
    "offboarding, basic network monitoring, and I write simple guides so common issues "
    "get solved faster next time.",
    "",
    "I believe this fits USIU-Africa well. A university needs a service desk that "
    "responds, systems that stay up during busy periods, classrooms and labs that are "
    "ready for teaching, and careful handling of staff and student data. I am comfortable "
    "with Windows, Microsoft 365, basic networking (TCP/IP, DNS, DHCP, VLANs), ticketing, "
    "asset tracking, patching, and MFA. I learn new systems quickly and I explain technical "
    "things in plain language.",
    "",
    "I hold a BSc in Applied Computing from KCA University (2023), plus certifications "
    "including CCNA and Microsoft Azure Fundamentals (AZ-900). I am ready to relocate to "
    "Nairobi and join your ICT team.",
    "",
    "Please find my CV, academic documents, and referees attached. Thank you for your time. "
    "I would welcome the chance to contribute to reliable ICT support for staff and students "
    "at USIU-Africa.",
    "",
    "Yours faithfully,",
    "Meshack Ariri",
    "+254 741 174 779",
    "aririmeshack@gmail.com",
    "linkedin.com/in/meshack-ariri",
]


def build_cover_letter(path):
    doc = Document()
    br.set_base_style(doc)
    br.add_header(doc, "MESHACK ARIRI", "ICT Officer — USIU-Africa", CONTACT_NAIROBI)
    for para in COVER_BODY:
        if para == "":
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(para)
        r.font.size = Pt(10.5)
        r.font.name = br.FONT_NAME
    doc.save(path)
    print("Saved:", path)


def export_copies(*paths):
    targets = [
        os.path.join(ROOT, "downloads", "JOBS", "USIU Africa"),
        os.path.join("/workspace", "JOBS", "USIU Africa"),
        "/opt/cursor/artifacts/meshack-cvs/USIU Africa",
        "/home/ccc/Documents/BASE/JOBS/USIU Africa",
    ]
    for dest in targets:
        try:
            os.makedirs(dest, exist_ok=True)
        except PermissionError:
            # Fall back with elevated copy handled outside if needed
            continue
        for src in paths:
            if os.path.isfile(src):
                shutil.copy2(src, dest)
                print("Copied:", os.path.join(dest, os.path.basename(src)))


if __name__ == "__main__":
    cv_docx = dict(cv)
    cv_docx["referees"] = cv_docx_referees
    cv_pdf = dict(cv)
    cv_pdf["referees"] = cv_docx_referees

    # Human-readable filenames (no SCREAMING_SNAKE)
    cv_docx_path = os.path.join(HERE, "Meshack Ariri CV.docx")
    cv_pdf_path = os.path.join(HERE, "Meshack Ariri CV.pdf")
    cl_docx_path = os.path.join(HERE, "Meshack Ariri Cover Letter.docx")
    cl_pdf_path = os.path.join(HERE, "Meshack Ariri Cover Letter.pdf")

    original_contact = br.CONTACT
    br.CONTACT = CONTACT_NAIROBI
    try:
        br.build(cv_docx, cv_docx_path)
        build_cover_letter(cl_docx_path)
        pb.build_cv_pdf(cv_pdf, cv_pdf_path, CONTACT_NAIROBI)
        pb.build_cover_pdf(
            COVER_BODY, cl_pdf_path, CONTACT_NAIROBI, "ICT Officer — USIU-Africa"
        )
    finally:
        br.CONTACT = original_contact

    export_copies(cv_docx_path, cv_pdf_path, cl_docx_path, cl_pdf_path)
    print("USIU-Africa ICT Officer application docs generated (DOCX + PDF).")
