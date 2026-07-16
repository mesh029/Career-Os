#!/usr/bin/env python3
"""Generate tailored CV + Cover Letter (.docx) for the ICT Manager (Hospitality) role.
Reuses styling from resumes/build_resumes.py. Facts stay consistent with 00_Master_Resume.md.
Run: python3 build.py
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
RESUMES = os.path.abspath(os.path.join(HERE, "..", "..", "resumes"))
sys.path.insert(0, RESUMES)
import build_resumes as br  # reuse CV builder + styling

# ---------------- Tailored CV data ----------------
cv = {
"headline": "ICT Manager  |  IT Infrastructure, Networks & Security  |  Systems Administration  |  Technical Support & Vendor Management",
"summary": ("IT professional with around three years at PATH running and supporting IT infrastructure for 120+ users "
    "across several sites, often as the main point of contact. Experienced across networks (LAN/WAN, VLANs, guest and "
    "staff Wi-Fi), firewalls (FortiGate, Cisco Meraki, MikroTik), Windows and Linux servers, virtualization (VMware, "
    "Proxmox), Microsoft 365 and Azure, backups and disaster recovery, and day to day technical support at around 98% "
    "uptime. Comfortable with IP telephony (3CX), CCTV, access control, and conference-room AV, and used to on-call work "
    "for systems that cannot go down. Has mentored junior staff, led network projects from planning to handover, managed "
    "a program IT budget, and coordinated ISPs and vendors. Keen to take on an ICT Manager role and get up to speed "
    "quickly on hospitality systems such as PMS, POS, and keycard platforms."),
"skills": [
    ("IT Infrastructure & Networks", "LAN/WAN, VLANs, VPN, guest and staff Wi-Fi (multi-AP), DNS/DHCP, routers and switches, Cisco Meraki, MikroTik, structured cabling, network monitoring"),
    ("Servers, Cloud & Virtualization", "Windows/Linux Server, Active Directory, Microsoft 365, Microsoft Azure, VMware and Proxmox, server-room management, backup and disaster recovery (Veeam/rsync)"),
    ("Cybersecurity & Data Protection", "Firewalls (FortiGate, Cisco Meraki MX, MikroTik), endpoint protection, MFA and identity/access management, patch management, data protection, security-awareness training"),
    ("Hospitality & Facilities Systems", "IP telephony/VoIP (3CX), CCTV and surveillance, access control, AV and conference-room setups; fast to learn PMS, POS and keycard systems"),
    ("Support & Service Management", "1st/2nd-line support for 120+ users, ITIL, SLA, incident and root-cause analysis, ticketing (Freshdesk/Zoho/Jira), on-call and after-hours support"),
    ("Leadership, Vendor & Assets", "Team mentoring and training, sole IT point of contact, project delivery, vendor/ISP coordination, IT budgeting, hardware and software-license inventory"),
],
"experience": [
    {"org":"PATH","title":"ICT / Health Informatics Officer","dates":"Mar 2025 to Present","loc":"Kenya (Hybrid)","bullets":[
        "Serve as the main IT point of contact for the program, keeping servers (Windows and Linux) and core systems running at over 98% uptime with backups, patching, encryption, and disaster-recovery plans.",
        "Manage networks end to end, including VLAN segmentation, firewalls (FortiGate, Cisco Meraki MX, MikroTik), VPN, and guest and staff Wi-Fi across multiple access points.",
        "Run virtualized infrastructure on VMware and Proxmox alongside Microsoft Azure, Microsoft 365, and Active Directory, with MFA and endpoint protection enforced throughout.",
        "Look after IT asset and software-license inventory, access controls, IP telephony (3CX), CCTV, and conference-room AV, and provide on-call support for systems that cannot go down.",
        "Proposed and managed a program IT budget, sourced quotes, and coordinated ISPs and external vendors.",
        "Designed and built a full platform (Next.js, Prisma, MySQL) for facility reporting, IT asset registers, and support tickets across four counties, and mentor junior IT staff.",
    ]},
    {"org":"PATH","title":"ICT Associate","dates":"Jan 2024 to Mar 2025","loc":"Kisumu, Kenya","bullets":[
        "Delivered 1st and 2nd-line support to 120+ users across sites, reaching around 98% uptime and about 25% faster response times under SLA.",
        "Configured and monitored Cisco Meraki and MikroTik networks (access points, switches, security appliances) and troubleshot LAN/WAN, VPN, Wi-Fi, DNS, and DHCP.",
        "Administered Active Directory and Microsoft 365 and enforced least-privilege access, group policy, MFA, and endpoint protection.",
        "Set up and maintained IP telephony (3CX), CCTV, and conference-room AV for staff and meetings.",
        "Led an office network setup and migration from planning and cabling through configuration and handover.",
        "Coordinated with ISPs and vendors, kept the IT asset inventory and knowledge base current, and mentored interns.",
    ]},
    {"org":"PATH","title":"ICT Intern","dates":"Apr 2023 to Dec 2023","loc":"Homa Bay, Kenya","bullets":[
        "Supported LAN/WAN infrastructure, routers, switches, structured cabling, printers, and VoIP, and assisted with server-room diagnostics.",
        "Logged and tracked incidents in Freshdesk under ITIL and SLA workflows and contributed to IT documentation and audits.",
    ]},
],
"certs": "CCNA  |  Microsoft Azure Fundamentals (AZ-900)  |  Google IT Support  |  IBM Technical Support  |  GitHub Professional  |  Microsoft Data Analysis",
"achievements": [
    "Kept mission-critical systems above 98% uptime across multi-site operations.",
    "Led an office network setup and migration from planning through to handover.",
    "Built a full platform managing facility reporting, IT asset registers, and incident tracking across four counties.",
    "Trained 80+ staff on secure system use, improving adoption and reducing recurring issues.",
    "Competed in cybersecurity hackathons, including the Cyberise Hackathon (Communications Authority of Kenya) and the NRF Cybersecurity Hackathon.",
],
}

# ---------------- Cover letter ----------------
COVER_BODY = [
    "Dear Hiring Manager,",
    "",
    "I am writing to apply for the ICT Manager position at your hotel. Over the past three years at PATH I have run and supported IT infrastructure for more than 120 users across several sites, often as the main point of contact, and I am strong in the exact areas this role calls for: networks, servers, cloud and virtualization, security, day to day technical support, IP telephony, CCTV and access control, and vendor and asset management. I have also mentored junior staff, led network projects end to end, and managed a program IT budget.",
    "In my current role I keep servers and core systems running at over 98% uptime through regular backups, patching, and recovery plans. I manage networks including VLANs, firewalls (FortiGate, Cisco Meraki, MikroTik), VPN, and guest and staff Wi-Fi, run virtualization on VMware and Proxmox alongside Microsoft 365 and Azure, and enforce MFA and endpoint protection. I also look after IP telephony (3CX), CCTV, access control, and conference room AV, coordinate ISPs and vendors, and provide on-call support for systems that cannot go down.",
    "I also build systems, not just fix them. I designed and built a full platform (Next.js, Prisma, MySQL) that tracks facility reporting, IT assets, and support tickets across four counties. It shows that I can take ownership of technology, plan improvements, and turn everyday operational needs into solutions that last.",
    "A good part of this role overlaps with what I already do, from networks and security to IP telephony, CCTV, and access control. Where I have less direct experience is the hotel specific platforms such as Property Management Systems, POS, and keycard systems, but I pick up new tools quickly and I am confident I can master them, especially given my networking and security background. I hold a BSc in Applied Computing, where I specialized in cybersecurity and digital forensics, along with CCNA and Microsoft Azure Fundamentals (AZ-900), and I have taken part in cybersecurity hackathons that keep me sharp on protecting systems and data.",
    "I would welcome the chance to talk about how my experience with infrastructure, security, and support can help keep your hotel's technology reliable, secure, and focused on the guest. Thank you for your time and consideration.",
    "",
    "Sincerely,",
    "Meshack Ariri",
]

def build_cover_letter(path):
    doc = Document()
    br.set_base_style(doc)
    br.add_name(doc, "MESHACK ARIRI")
    br.add_title(doc, "ICT Manager")
    br.add_contact(doc, br.CONTACT)
    for para in COVER_BODY:
        if para == "":
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(para)
        r.font.size = Pt(10.5)
    doc.save(path)
    print("Saved:", path)


if __name__ == "__main__":
    import pdf_builder as pb
    # DOCX
    br.build(cv, os.path.join(HERE, "Meshack_Ariri_CV_ICT_Manager.docx"))
    build_cover_letter(os.path.join(HERE, "Meshack_Ariri_Cover_Letter_ICT_Manager.docx"))
    # PDF (themed, deep teal)
    pb.build_cv_pdf(cv, os.path.join(HERE, "Meshack_Ariri_CV_ICT_Manager.pdf"), br.CONTACT)
    pb.build_cover_pdf(COVER_BODY, os.path.join(HERE, "Meshack_Ariri_Cover_Letter_ICT_Manager.pdf"), br.CONTACT, "ICT Manager")
    print("ICT Manager application docs generated (DOCX + PDF).")
