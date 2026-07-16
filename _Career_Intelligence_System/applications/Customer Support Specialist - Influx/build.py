#!/usr/bin/env python3
"""Tailored CV + Cover Letter (DOCX + PDF) for Customer Support Specialist, Influx.
Facts per 00_Master_Resume.md. Human tone, no em-dashes. Run: python3 build.py
"""
import os, sys
from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
RESUMES = os.path.abspath(os.path.join(HERE, "..", "..", "resumes"))
sys.path.insert(0, RESUMES)
import build_resumes as br
import pdf_builder as pb

cv = {
"headline": "Customer Support Specialist  |  Technical Helpdesk & User Support  |  Freshdesk · ITIL · Remote-ready",
"summary": ("Helpdesk and user support professional with about three years supporting 120+ users at PATH "
    "through email-style tickets, remote sessions, and in-person assistance. Resolves login, password, and "
    "account access issues, basic software setup, and common connectivity problems using clear step-by-step "
    "guidance. Logs and categorises cases in Freshdesk and Zoho Desk under ITIL and SLA, writes knowledge-base "
    "articles, and escalates non-standard issues with accurate notes. Improved average response time by about "
    "25% through better documentation and workflows. Patient communicator, strong in English, and comfortable "
    "learning new platforms quickly in a remote setup."),
"skills": [
    ("Customer & Helpdesk Support", "User-facing support, email and ticket-based assistance, live troubleshooting sessions, calm and empathetic communication, explaining technical steps in plain language"),
    ("Technical Troubleshooting", "Login and password issues, account access, MFA enrollment, Microsoft 365 and Teams, browser and connectivity problems, basic software setup and configuration, device onboarding"),
    ("Ticketing & CRM Tools", "Freshdesk, Zoho Desk, ticket categorisation, SLA tracking, escalation and handover notes, knowledge base articles, ITIL service desk workflows"),
    ("Documentation & Process", "Accurate case notes, recurring-issue feedback, SOPs and job aids, service metrics, process improvements that cut response time ~25%"),
    ("Platforms & Learning", "Web applications, Microsoft 365, mobile and desktop apps, SaaS onboarding concepts (HubSpot/Salesforce familiarity), quick ramp on new client tools"),
    ("Remote Work Setup", "Windows 10/11, reliable home internet, wired headset, quiet workspace, hybrid and remote delivery experience"),
],
"experience": [
    {"org":"PATH","title":"ICT / User Support Officer","dates":"Mar 2025 to Present","loc":"Kenya (Hybrid / Remote)","bullets":[
        "Support users with account access, passwords, MFA, and Microsoft 365 issues, explaining fixes clearly and documenting each case.",
        "Handle escalated helpdesk tickets that need deeper troubleshooting before resolution or specialist handoff.",
        "Maintain knowledge-base articles and short guides so common questions are answered faster on the next ticket.",
        "Train facility and program staff (80+ people across 15+ sessions) using patient, step-by-step communication.",
    ]},
    {"org":"PATH","title":"ICT Associate","dates":"Jan 2024 to Mar 2025","loc":"Kisumu, Kenya","bullets":[
        "Delivered 1st and 2nd-line user support to 120+ people via Freshdesk and Zoho Desk, often handling 20+ tickets per day at busy periods.",
        "Resolved login, password, account access, printer, VoIP, and basic connectivity issues following ITIL workflows and SLA targets.",
        "Walked users through software setup and configuration (Microsoft 365, Teams, SharePoint, devices) until the issue was closed.",
        "Categorised tickets accurately, wrote handover notes for escalations, and kept the knowledge base up to date.",
        "Reviewed recurring ticket types and suggested workflow tweaks that improved average response time by about 25%.",
        "Escalated complex network or server issues to senior staff with clear diagnostics of what was already tried.",
    ]},
    {"org":"PATH","title":"ICT Intern","dates":"Apr 2023 to Dec 2023","loc":"Homa Bay, Kenya","bullets":[
        "Logged and tracked user issues in Freshdesk, supported desktop and printer setup, and contributed to support documentation.",
    ]},
    {"org":"Independent","title":"Remote SaaS & IT Support (Project-Based)","dates":"2023 to Present","loc":"Remote","bullets":[
        "Provided remote setup and user onboarding for small businesses on Microsoft 365, Zoho, and Google Workspace.",
    ]},
],
"certs": "Google IT Support  |  IBM Technical Support  |  Microsoft Azure Fundamentals (AZ-900)  |  CCNA  |  GitHub Professional",
"achievements": [
    "Supported 120+ users with helpdesk-style technical support at over 98% system availability.",
    "Improved average ticket response time by about 25% through documentation and process review.",
    "Authored knowledge-base content and trained 80+ users on secure, practical system use.",
    "Handled high ticket volume with accurate logging, categorisation, and escalation when required.",
],
}

COVER_BODY = [
    "Dear Hiring Team,",
    "",
    "I am applying for the Customer Support Specialist role at Influx. For the past three years at PATH I have supported more than 120 users through helpdesk-style work: logging and resolving tickets in Freshdesk, helping with login and password problems, Microsoft 365 and account access, basic device and software setup, and connectivity issues that users could see on their side. I explain fixes in plain language, document every interaction properly, and escalate when something falls outside the standard workflow.",
    "I improved our average ticket response time by about 25% by keeping the knowledge base current and tightening how we categorise and hand off cases. I have also trained 80+ colleagues on secure system use, which helped me practice patient, step-by-step communication with people who are not technical.",
    "I am fluent in English, comfortable learning new tools quickly, and genuinely enjoy the moment a user understands the fix. I am looking for a full-time remote role with an international team, and Influx fits that well.",
    "I work from a quiet home setup on Windows with 8GB+ RAM, Intel Core i5 or better, wired headset, and reliable broadband (40+ Mbps). I would welcome the chance to contribute to your global support team.",
    "",
    "Thank you for considering my application.",
    "",
    "Sincerely,",
    "Meshack Ariri",
]

def build_cover_letter(path):
    doc = Document()
    br.set_base_style(doc)
    br.add_header(doc, "MESHACK ARIRI", "Customer Support Specialist", br.CONTACT)
    for para in COVER_BODY:
        if para == "":
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(para)
        r.font.size = Pt(10.5)
    doc.save(path)
    print("Saved:", path)

if __name__ == "__main__":
    base = "Meshack_Ariri"
    br.build(cv, os.path.join(HERE, f"{base}_CV_Customer_Support_Influx.docx"))
    build_cover_letter(os.path.join(HERE, f"{base}_Cover_Letter_Customer_Support_Influx.docx"))
    pb.build_cv_pdf(cv, os.path.join(HERE, f"{base}_CV_Customer_Support_Influx.pdf"), br.CONTACT)
    pb.build_cover_pdf(
        COVER_BODY,
        os.path.join(HERE, f"{base}_Cover_Letter_Customer_Support_Influx.pdf"),
        br.CONTACT,
        "Customer Support Specialist",
    )
    print("Influx application docs generated (DOCX + PDF).")
