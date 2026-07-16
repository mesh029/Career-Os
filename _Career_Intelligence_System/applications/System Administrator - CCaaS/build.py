#!/usr/bin/env python3
"""Tailored CV + Cover Letter (DOCX + PDF) for System Administrator (CCaaS & Enterprise Stack).
Achievement-based positioning per Application_Intelligence.md. Facts per 00_Master_Resume.md.
Run: python3 build.py
"""
import os, sys
from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
RESUMES = os.path.abspath(os.path.join(HERE, "..", "..", "resumes"))
sys.path.insert(0, RESUMES)
import build_resumes as br
import pdf_builder as pb

cv = {
"headline": "Systems Administrator  |  Enterprise SaaS & CCaaS Stack  |  IT Auditing, License Management & Reporting",
"summary": ("Systems administrator with about three years running IT operations for 120+ users across multiple "
    "sites. Focused on the strategic side of the stack: auditing which tools are actually being used, keeping "
    "licenses and access right-sized, and getting telephony, ITSM, identity (Active Directory and Entra ID with "
    "MFA), and reporting to work as one system rather than in silos. Comfortable building the reports leadership "
    "relies on for renewal and budget decisions (SQL, Power BI, Grafana), and designed and built a full "
    "operations platform with 40+ REST APIs for IT asset tracking, incident ticketing, data validation, and "
    "role-based access. Works independently, owns projects from business case through rollout, and prefers "
    "backend ownership over day-to-day helpdesk."),
"skills": [
    ("CCaaS & Contact-Center Stack", "VoIP/3CX, call routing, ITSM integration (Freshdesk, Zoho Desk), Microsoft Teams, SLA and queue handling; CCaaS building blocks (QA, reporting, SSO, security)"),
    ("IT Auditing & License Management", "Software stack audits, SaaS license utilization, access reviews, deprovisioning, hardware and software asset registers, reducing wasted spend"),
    ("Identity, SSO & Security", "Microsoft 365, Entra ID/Azure AD, Active Directory, MFA, conditional access, Group Policy, least-privilege access, endpoint protection, patch management"),
    ("Reporting & Analytics", "Power BI, Grafana, Metabase, SQL (MySQL/PostgreSQL), operational dashboards, utilization and audit reports, presenting to leadership"),
    ("Enterprise Tool Framework", "SOPs, onboarding and offboarding playbooks, knowledge base, ITIL workflows, vendor evaluation, contract renewal prep, API and system integration"),
    ("Automation & Platforms", "PowerShell, Bash, Python for provisioning and reporting; REST APIs and middleware; Next.js, TypeScript, Prisma, MySQL"),
],
"experience": [
    {"org":"PATH","title":"ICT / Systems Officer (HMIS)","dates":"Mar 2025 to Present","loc":"Kenya (Hybrid)","bullets":[
        "Keep mission-critical Windows and Linux application servers above 98% uptime through patching, encryption, backups, and disaster recovery.",
        "Run regular system and data-flow audits that surface access gaps and integration risks, then turn the findings into SOPs and concrete fixes.",
        "Build the operational reporting leadership uses (Power BI, Grafana, SQL), including automated checks that cut manual reporting time by about 40%.",
        "Manage identity and security across Microsoft 365 and Active Directory (MFA, least-privilege access) for 120+ users across several sites.",
    ]},
    {"org":"PATH","title":"ICT Associate","dates":"Jan 2024 to Mar 2025","loc":"Kisumu, Kenya","bullets":[
        "Ran IT operations for 120+ users across multiple sites at around 98% uptime, administering Microsoft 365, Active Directory and Entra ID, and MFA-based single sign-on with least-privilege access.",
        "Set up and ran the program's contact-centre-adjacent stack: VoIP/3CX telephony connected to Freshdesk and Zoho Desk, with ITIL workflows and SLA tracking that cut average ticket response time by about 25%.",
        "Ran regular software and license audits across Microsoft 365, endpoints, and SaaS accounts, removed inactive users, right-sized access, and kept a live asset and license inventory to stop wasted spend.",
        "Designed how new staff are onboarded and offboarded (account setup, devices, Microsoft 365 and Teams access, SOPs, knowledge-base articles) so operations and support teams could handle the frontline steps without waiting on central IT.",
        "Handled ISP and vendor evaluations, gathered requirements and quotes, and prepared renewal recommendations for leadership sign-off on software and connectivity contracts.",
        "Configured and monitored Cisco Meraki network equipment for reliable, secure connectivity across sites.",
    ]},
    {"org":"Independent","title":"Enterprise Systems Platform (Project-Based)","dates":"2023 to Present","loc":"Remote","bullets":[
        "Designed and built a full IT operations platform (Next.js, TypeScript, Prisma, MySQL on Aiven) with more than 40 REST APIs, covering IT asset tracking, incident and NOC ticketing, data validation, reporting dashboards, and role-based access.",
        "Took it from idea to working product: business case, build, integration, documentation, and the utilization reporting leadership would use (built with demo data across four counties).",
    ]},
    {"org":"PATH","title":"ICT Intern","dates":"Apr 2023 to Dec 2023","loc":"Homa Bay, Kenya","bullets":[
        "Helped roll out multi-site IT infrastructure (LAN/WAN, VoIP, PXE desktop imaging, structured cabling) and kept the Freshdesk ticket and SLA records that later became the basis for the program's ITIL service-management setup.",
    ]},
],
"certs": "Microsoft Azure Fundamentals (AZ-900)  |  CCNA  |  Google IT Support  |  IBM Technical Support  |  GitHub Professional  |  Microsoft Data Analysis",
"achievements": [
    "Ran an integrated stack (VoIP/3CX, ITSM, Microsoft 365, Active Directory SSO) for 120+ users at over 98% uptime.",
    "Cut manual reporting and processing time by about 40% with SQL automation and dashboards leadership actually uses.",
    "Built an operations platform with 40+ REST APIs covering asset tracking, ticketing, data validation, and analytics.",
    "Set up onboarding processes and trained 80+ staff across 15+ sessions on secure system use.",
    "Improved average ticket response time by about 25% by tightening ITIL and SLA workflows.",
],
}

COVER_BODY = [
    "Dear Hiring Manager,",
    "",
    "I am applying for the System Administrator role focused on CCaaS and enterprise stack optimization. This is the kind of work I already do: figuring out which tools are actually earning their keep, keeping licenses and access right-sized, getting telephony and ITSM to run as one system, and giving leadership reporting they can act on rather than chasing tickets all day.",
    "At PATH I have run IT operations for more than 120 users across several sites at around 98% uptime, administering Microsoft 365, Entra ID and Active Directory (with MFA and least-privilege), VoIP/3CX, and Freshdesk and Zoho Desk under ITIL and SLA. I run regular software and access audits, keep hardware and license inventories current, and prepare the utilization summaries that inform renewal and vendor decisions. I also built automated reporting (SQL, Power BI, Grafana) that cut manual processing time by about 40%.",
    "What I bring beyond standard administration is that I build the frameworks, not just maintain them. I designed and built a full operations platform (Next.js, Prisma, MySQL, 40+ REST APIs) with modules for IT asset tracking, incident and NOC ticketing, data validation, and executive dashboards, which lines up closely with your need for a solid way to deploy and manage enterprise tools.",
    "I want to be straightforward about one thing: my deepest hands-on experience is with VoIP/3CX and ITSM tied into identity and analytics, rather than a named CCaaS platform like Five9 or Genesys. The architecture, the auditing habits, the SSO and security work, and the reporting all carry over directly, and I am working through an Amazon Connect sandbox to close that gap quickly.",
    "I work independently, keep communication clear with leadership, and take projects from the first business case through to rollout and audit. I would welcome the chance to take full ownership of your CCaaS environment and wider software stack.",
    "",
    "Thank you for your time and consideration.",
    "",
    "Sincerely,",
    "Meshack Ariri",
]

def build_cover_letter(path):
    doc = Document()
    br.set_base_style(doc)
    br.add_header(doc, "MESHACK ARIRI", "System Administrator", br.CONTACT)
    for para in COVER_BODY:
        if para == "":
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(para)
        r.font.size = Pt(10.5)
    doc.save(path)
    print("Saved:", path)

if __name__ == "__main__":
    br.build(cv, os.path.join(HERE, "Meshack_Ariri_CV_System_Administrator_CCaaS.docx"))
    build_cover_letter(os.path.join(HERE, "Meshack_Ariri_Cover_Letter_System_Administrator_CCaaS.docx"))
    pb.build_cv_pdf(cv, os.path.join(HERE, "Meshack_Ariri_CV_System_Administrator_CCaaS.pdf"), br.CONTACT)
    pb.build_cover_pdf(
        COVER_BODY,
        os.path.join(HERE, "Meshack_Ariri_Cover_Letter_System_Administrator_CCaaS.pdf"),
        br.CONTACT,
        "System Administrator",
    )
    print("CCaaS System Administrator application docs generated (DOCX + PDF).")
