#!/usr/bin/env python3
"""Tailored CV + Cover Letter (DOCX + PDF) for IT Officer – Global Support (INSO).
Reuses resumes/build_resumes.py (DOCX) and resumes/pdf_builder.py (PDF). Facts per 00_Master_Resume.md.
Run: python3 build.py
"""
import os, sys
from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
RESUMES = os.path.abspath(os.path.join(HERE, "..", "..", "resumes"))
sys.path.insert(0, RESUMES)
import build_resumes as br
import pdf_builder as pb

cv = {
"headline": "IT Officer – Global Support  |  Humanitarian / INGO IT  |  Microsoft 365, Entra ID & Intune  |  Remote & Field Support",
"summary": ("IT professional with around three years at PATH, a global health NGO, delivering 1st- and 2nd-line "
    "support for 120+ users across multiple sites at roughly 98% uptime. Experienced with Windows, Microsoft 365 "
    "(Exchange Online, Teams, SharePoint Online), Entra ID / Active Directory, and Intune-style endpoint deployment, "
    "plus remote troubleshooting via Microsoft Teams, cybersecurity awareness training, network and Wi-Fi diagnostics, "
    "IT onboarding, SharePoint permissions guidance, documentation, and coordinating support between field offices "
    "and central IT. Legally authorised to work in Kenya; open to frequent travel for project-based mission support."),
"skills": [
    ("Microsoft 365 & Cloud", "Entra ID / Azure AD, Exchange Online, Teams, SharePoint Online, MFA, conditional access, user provisioning, Microsoft Azure (AZ-900)"),
    ("Endpoint & Intune", "Intune / MDM device deployment and management, Windows 10/11 imaging and configuration, patching, endpoint protection, onboarding/offboarding"),
    ("IT Support & Remote Delivery", "1st/2nd-line support, Microsoft Teams remote support, hardware/software troubleshooting, IT onboarding, TeamViewer-style remote assistance"),
    ("Networking", "LAN/WAN, Wi-Fi, DNS/DHCP, VPN, TCP/IP, Cisco Meraki, MikroTik; CCNA"),
    ("Cybersecurity & Awareness", "MFA, endpoint protection, patch management, access controls, security training and awareness materials; cybersecurity & digital forensics foundation"),
    ("INGO Operations", "Multi-site field support, SOPs and tutorials, knowledge base, IT asset tracking, cross-cultural communication, ITIL/SLA (Freshdesk, Zoho Desk)"),
],
"experience": [
    {"org":"PATH","title":"Health Informatics / HMIS Officer","dates":"Mar 2025 to Present","loc":"Kisumu/Nyamira, Kenya (Hybrid)","bullets":[
        "Provide IT support alongside HMIS duties for 120+ users across sites, sustaining over 98% uptime for mission-critical systems.",
        "Administer Microsoft 365, Entra ID / Active Directory, and Intune-style endpoint management; support SharePoint and Teams.",
        "Deliver remote and in-person troubleshooting, IT onboarding, and user training with clear documentation and SOPs.",
        "Troubleshoot network and Wi-Fi connectivity; maintain IT asset records and coordinate with vendors and field teams.",
    ]},
    {"org":"PATH","title":"ICT Associate","dates":"Jan 2024 to Mar 2025","loc":"Kisumu, Kenya","bullets":[
        "Delivered 1st- and 2nd-line IT support to 120+ users across multiple sites, reaching around 98% uptime and about 25% faster response under SLA.",
        "Installed and configured Windows workstations, Microsoft 365 applications, and network settings to organisational standards.",
        "Administered Active Directory and Microsoft 365 accounts; supported SharePoint, Teams, MFA, and endpoint protection.",
        "Managed endpoints through Intune-style deployment; troubleshot LAN/WAN, Wi-Fi, DNS, and VPN on Cisco Meraki and MikroTik networks.",
        "Authored knowledge-base articles, delivered onboarding sessions, and logged incidents in Freshdesk and Zoho Desk under ITIL/SLA.",
        "Maintained IT asset inventory and supported equipment deployment across field locations.",
    ]},
    {"org":"PATH","title":"ICT Intern","dates":"Apr 2023 to Dec 2023","loc":"Homa Bay, Kenya","bullets":[
        "Supported LAN/WAN infrastructure, Wi-Fi, printers, VoIP, and desktop imaging; assisted server-room diagnostics.",
        "Logged and tracked incidents in Freshdesk; contributed to IT documentation and field deployment support.",
    ]},
],
"certs": "CCNA  |  Microsoft Azure Fundamentals (AZ-900)  |  Google IT Support  |  IBM Technical Support  |  GitHub Professional  |  Microsoft Data Analysis",
"achievements": [
    "Supported 120+ NGO staff at over 98% uptime across multi-site humanitarian health operations.",
    "Delivered IT onboarding, training, and documentation for field and office teams.",
    "Administered Microsoft 365 and Entra ID with MFA, SharePoint, and endpoint management.",
    "Competed in cybersecurity hackathons (Cyberise / Communications Authority of Kenya, NRF).",
],
"languages": "English (professional working proficiency)  |  French (basic — actively improving)",
}

COVER_BODY = [
    "Dear Hiring Team,",
    "",
    "I am applying for the IT Officer – Global Support role at INSO. Over the past three years at PATH, a global health NGO, I have provided 1st- and 2nd-line IT support for more than 120 users across multiple sites, keeping mission-critical systems running at over 98% uptime while supporting field teams in fast-moving, multi-location environments.",
    "My daily work aligns closely with this role: administering Microsoft 365, Entra ID, Exchange Online, Teams, and SharePoint Online; deploying and managing endpoints through Intune-style device management; troubleshooting hardware, software, and network issues remotely and in person; delivering IT onboarding and training; maintaining accurate documentation and asset records; and acting as a coordination point between field offices and central IT. I also contribute to cybersecurity awareness through MFA enforcement, endpoint protection, patching, and user guidance.",
    "I am legally authorised to work in Kenya and willing to travel frequently to INSO country missions for project-based support. My professional working language is English; I have basic French knowledge and I am actively improving it. I communicate technical concepts clearly to non-technical staff across cultures and work effectively in distributed, remote teams.",
    "INSO's mission to protect humanitarian workers in insecure contexts is work I would be proud to support. I would welcome the opportunity to help standardise and strengthen IT across your country offices. Thank you for your consideration.",
    "",
    "Sincerely,",
    "Meshack Ariri",
]

def build_cover_letter(path):
    doc = Document()
    br.set_base_style(doc)
    br.add_name(doc, "MESHACK ARIRI")
    br.add_title(doc, "IT Officer – Global Support")
    br.add_contact(doc, br.CONTACT)
    for para in COVER_BODY:
        if para == "":
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(para)
        r.font.size = Pt(10.5)
    doc.save(path)
    print("Saved:", path)

if __name__ == "__main__":
    br.build(cv, os.path.join(HERE, "Meshack_Ariri_CV_IT_Officer_Global_Support_INSO.docx"))
    build_cover_letter(os.path.join(HERE, "Meshack_Ariri_Cover_Letter_IT_Officer_Global_Support_INSO.docx"))
    pb.build_cv_pdf(cv, os.path.join(HERE, "Meshack_Ariri_CV_IT_Officer_Global_Support_INSO.pdf"), br.CONTACT)
    pb.build_cover_pdf(COVER_BODY, os.path.join(HERE, "Meshack_Ariri_Cover_Letter_IT_Officer_Global_Support_INSO.pdf"), br.CONTACT, "IT Officer – Global Support")
    print("INSO application docs generated (DOCX + PDF).")
