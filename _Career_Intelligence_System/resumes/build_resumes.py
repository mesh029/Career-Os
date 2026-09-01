#!/usr/bin/env python3
"""
Career Intelligence System - Resume builder.
Generates ATS-optimized .docx resumes from structured data (single source of truth: 00_Master_Resume.md).
Design: modern header band, accent-bar section headings, single-column, standard fonts, scannable < 30s.
Run:  python3 build_resumes.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from resume_theme import (
    ACCENT_HEX, ACCENT_DARK_HEX, ACCENT_MID_HEX, ACCENT_LIGHT_HEX, ACCENT_PALE_HEX,
    SLATE_HEX, SLATE_MUTED_HEX, FONT, CONTACT, format_contact_lines,
)


def _rgb(hex6):
    return RGBColor(int(hex6[0:2], 16), int(hex6[2:4], 16), int(hex6[4:6], 16))


ACCENT = _rgb(ACCENT_HEX)
ACCENT_DARK = _rgb(ACCENT_DARK_HEX)
ACCENT_MID = _rgb(ACCENT_MID_HEX)
DARK = _rgb(SLATE_HEX)
GREY = _rgb(SLATE_MUTED_HEX)
FONT_NAME = FONT

CONTACT = CONTACT  # re-export for application build scripts


def _set_cell_shading(cell, fill_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill_hex)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(shading)


def _remove_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)


def _set_table_width(table, width_inches):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)


def set_base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT_NAME
    st.font.size = Pt(10.5)
    st.font.color.rgb = DARK
    pf = st.paragraph_format
    pf.space_after = Pt(2)
    pf.line_spacing = 1.05
    for section in doc.sections:
        section.top_margin = Inches(0.48)
        section.bottom_margin = Inches(0.48)
        section.left_margin = Inches(0.58)
        section.right_margin = Inches(0.58)


def _render_header_band(doc, name, headline, contact):
    """Full-width shaded header with top accent stripe."""
    content_w = Inches(6.84)

    # Top accent stripe
    stripe = doc.add_table(rows=1, cols=1)
    _set_table_width(stripe, 6.84)
    _remove_table_borders(stripe)
    sc = stripe.rows[0].cells[0]
    _set_cell_shading(sc, ACCENT_DARK_HEX)
    sp = sc.paragraphs[0]
    sp.paragraph_format.space_after = Pt(0)
    sp.paragraph_format.space_before = Pt(0)
    run = sp.add_run(" ")
    run.font.size = Pt(3)

    # Main header band
    band = doc.add_table(rows=1, cols=1)
    _set_table_width(band, 6.84)
    _remove_table_borders(band)
    cell = band.rows[0].cells[0]
    _set_cell_shading(cell, ACCENT_LIGHT_HEX)
    cell.vertical_alignment = 1  # center

    # Clear default paragraph
    cell.paragraphs[0].clear()

    p_name = cell.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(10)
    p_name.paragraph_format.space_after = Pt(4)
    rn = p_name.add_run(name)
    rn.bold = True
    rn.font.size = Pt(26)
    rn.font.color.rgb = ACCENT_DARK
    rn.font.name = FONT_NAME

    p_title = cell.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(6)
    rt = p_title.add_run(headline)
    rt.font.size = Pt(10.2)
    rt.font.color.rgb = DARK
    rt.font.name = FONT_NAME

    for line in format_contact_lines(contact):
        p_contact = cell.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_contact.paragraph_format.space_after = Pt(2)
        parts = [x.strip() for x in line.split("|")]
        for i, part in enumerate(parts):
            if i > 0:
                sep = p_contact.add_run("  |  ")
                sep.font.size = Pt(8.4)
                sep.font.color.rgb = ACCENT_MID
                sep.font.name = FONT_NAME
            rc = p_contact.add_run(part)
            rc.font.size = Pt(8.4)
            rc.font.color.rgb = GREY
            rc.font.name = FONT_NAME

    # padding bottom
    p_pad = cell.add_paragraph()
    p_pad.paragraph_format.space_after = Pt(10)

    # Rule under header
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(4)
    rule.paragraph_format.space_after = Pt(8)
    p_pr = rule._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "16")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT_HEX)
    pbdr.append(bottom)
    p_pr.append(pbdr)


def _init_header_state(doc):
    if not hasattr(doc, "_cv_header"):
        doc._cv_header = {"_rendered": False}


def _maybe_render_header(doc):
    h = doc._cv_header
    if h["_rendered"]:
        return
    if not all(k in h for k in ("name", "title", "contact")):
        return
    h["_rendered"] = True
    _render_header_band(doc, h["name"], h["title"], h["contact"])


def add_header(doc, name, headline, contact):
    """Render the full header band in one call."""
    _render_header_band(doc, name, headline, contact)
    doc._cv_header = {"_rendered": True}


def add_name(doc, name):
    _init_header_state(doc)
    doc._cv_header["name"] = name
    _maybe_render_header(doc)


def add_title(doc, title):
    _init_header_state(doc)
    doc._cv_header["title"] = title
    _maybe_render_header(doc)


def add_contact(doc, contact):
    _init_header_state(doc)
    doc._cv_header["contact"] = contact
    _maybe_render_header(doc)


def add_heading(doc, text):
    """Section heading with left accent bar on pale background."""
    bar_w = Inches(0.12)
    content_w = Inches(6.72)
    row = doc.add_table(rows=1, cols=2)
    _set_table_width(row, 6.84)
    _remove_table_borders(row)

    bar_cell = row.rows[0].cells[0]
    bar_cell.width = bar_w
    _set_cell_shading(bar_cell, ACCENT_HEX)
    bar_cell.paragraphs[0].clear()

    text_cell = row.rows[0].cells[1]
    text_cell.width = content_w
    _set_cell_shading(text_cell, ACCENT_PALE_HEX)
    text_cell.paragraphs[0].clear()
    p = text_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Inches(0.08)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(10.8)
    r.font.color.rgb = ACCENT_DARK
    r.font.name = FONT_NAME

    # Subtle rule under section heading
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after = Pt(6)
    p_pr = rule._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT_MID_HEX)
    pbdr.append(bottom)
    p_pr.append(pbdr)


def add_para(doc, text, size=10.5, space_after=3, bold=False, color=DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = color
    r.font.name = FONT_NAME
    return p


def add_role(doc, org, title, dates, location=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    left = p.add_run(title)
    left.bold = True
    left.font.size = Pt(10.3)
    left.font.color.rgb = ACCENT
    left.font.name = FONT_NAME
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.95), alignment=3)
    p.add_run("\t")
    d = p.add_run(dates)
    d.font.size = Pt(8.8)
    d.font.color.rgb = GREY
    d.font.name = FONT_NAME

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.left_indent = Inches(0.06)
    o = p2.add_run(org)
    o.italic = True
    o.font.size = Pt(9.1)
    o.font.color.rgb = GREY
    o.font.name = FONT_NAME
    if location:
        sep = p2.add_run("   |   ")
        sep.font.size = Pt(9.1)
        sep.font.color.rgb = ACCENT_MID
        sep.font.name = FONT_NAME
        loc = p2.add_run(location)
        loc.italic = True
        loc.font.size = Pt(9.1)
        loc.font.color.rgb = GREY
        loc.font.name = FONT_NAME


def add_bullets(doc, bullets):
    for b in bullets:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.08
        r0 = p.add_run("▪  ")
        r0.font.color.rgb = ACCENT
        r0.bold = True
        r0.font.name = FONT_NAME
        r = p.add_run(b)
        r.font.size = Pt(9.35)
        r.font.name = FONT_NAME
        r.font.color.rgb = DARK


def add_skill_line(doc, label, items):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.04)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.size = Pt(9.3)
    r.font.color.rgb = ACCENT
    r.font.name = FONT_NAME
    r2 = p.add_run(items)
    r2.font.size = Pt(9.3)
    r2.font.name = FONT_NAME
    r2.font.color.rgb = DARK


def build(data, path):
    doc = Document()
    set_base_style(doc)
    add_header(doc, "MESHACK ARIRI", data["headline"], CONTACT)

    add_heading(doc, "Professional Summary")
    add_para(doc, data["summary"], size=10.2, space_after=4)

    add_heading(doc, "Core Skills")
    for label, items in data["skills"]:
        add_skill_line(doc, label, items)

    add_heading(doc, "Professional Experience")
    for role in data["experience"]:
        add_role(doc, role["org"], role["title"], role["dates"], role.get("loc", ""))
        add_bullets(doc, role["bullets"])

    add_heading(doc, "Education")
    add_para(doc, "BSc, Applied Computing  |  KCA University, Nairobi  |  November 2023", bold=True, size=10.2, space_after=1)
    add_para(doc, "Specialized in Cybersecurity and Digital Forensics. Coursework included Ethical Hacking, Digital Forensics, Data Protection, and Cloud Computing.", size=9.8, color=GREY)
    add_para(doc, "Capstone: Designed and implemented a secure office network (VLANs, Active Directory, backup/DR) simulating real-world IT administration.", size=9.8, color=GREY)

    if data.get("languages"):
        add_heading(doc, "Languages")
        add_para(doc, data["languages"], size=10)

    add_heading(doc, "Certifications")
    add_para(doc, data["certs"], size=10)

    if data.get("achievements"):
        add_heading(doc, "Achievements & Activities")
        add_bullets(doc, data["achievements"])

    add_heading(doc, "Referees")
    if data.get("referees"):
        for ref in data["referees"]:
            add_para(doc, ref, size=10, space_after=4)
    else:
        add_para(doc, "Available upon request.", size=10, color=GREY)

    doc.save(path)
    print("Saved:", path)

# ---------------------------------------------------------------------------
# SHARED building blocks
CERTS_HEALTH = ("CCNA  |  Microsoft Azure Fundamentals (AZ-900)  |  Google IT Support  |  "
                "IBM Technical Support  |  GitHub Professional  |  Microsoft Data Analysis")

ACH = [
    "Improved multi-system data quality and sustained >98% uptime across supported program sites.",
    "Automated validation and reporting workflows, reducing reporting/processing time by ~40%.",
    "Trained 80+ health workers and facility teams on digital tools, data quality, and secure system use.",
    "Competed in cybersecurity hackathons, including the Cyberise Hackathon (Communications Authority of Kenya) and the NRF Cybersecurity Hackathon.",
]

# 1) PRIMARY — Health Informatics
health = {
"headline": "Health Informatics Specialist  |  EMR (KenyaEMR/OpenMRS) & DHIS2  |  Health Data & Systems Support",
"summary": ("Health Informatics and IT professional with ~3 years at PATH supporting national-scale health "
    "information systems such as KenyaEMR/TaifaCare, DHIS2, and EMR, laboratory, and pharmacy integrations. Combines hands-on "
    "EMR/HMIS support with data quality (DQA, SQL validation), analytics (Power BI, Grafana), and reliable IT "
    "operations for 120+ users across multiple sites. Proven at improving uptime, safeguarding data integrity, and "
    "training health workers to adopt digital tools. Focused on digital-health and EMR/DHIS2 implementation roles."),
"skills": [
    ("Health Information Systems", "KenyaEMR/TaifaCare, OpenMRS, DHIS2, NDWH, EMR-lab-pharmacy integration, interoperability, data quality (DQA), surveillance/MEL reporting, Kobo Toolbox"),
    ("Data & Analytics", "SQL (MySQL/PostgreSQL), Power BI, Grafana, Metabase, advanced Excel, Python (automation), data validation & pipelines"),
    ("IT & Systems", "Microsoft 365, Active Directory/Azure AD, Windows/Linux servers, Apache Tomcat, backups & disaster recovery, networking"),
    ("Service Management", "Freshdesk, Zoho Desk, Jira Service Management, ITIL, SLA, incident & root-cause analysis, knowledge base"),
    ("Delivery & Enablement", "Training & capacity building, technical documentation/SOPs, stakeholder engagement, requirements analysis"),
],
"experience": [
    {"org":"PATH","title":"Health Informatics / HMIS Officer","dates":"Mar 2025 – Present","loc":"Kisumu/Nyamira, Kenya (Hybrid)","bullets":[
        "Support and stabilize KenyaEMR/TaifaCare and DHIS2 across health facilities, sustaining >98% uptime for critical clinical services.",
        "Develop SQL validation and data-quality (DQA) workflows that improve early error detection and reduce reporting time by ~40%.",
        "Build Power BI, Grafana, and Metabase dashboards for surveillance, trend analysis, and program decision-making.",
        "Support EMR-laboratory-pharmacy integrations and multi-system data flow, strengthening interoperability and reporting accuracy.",
        "Maintain Windows and Linux servers hosting EMR data, covering backups, encryption, patching, and disaster recovery.",
        "Conduct system audits, data-flow reviews, and risk identification aligned to national HIS standards.",
        "Train facility staff, ICT officers, and program teams on EMR modules, DHIS2, data quality, and secure system use.",
    ]},
    {"org":"PATH","title":"ICT Associate","dates":"Jan 2024 – Mar 2025","loc":"Kisumu, Kenya","bullets":[
        "Provided 1st- and 2nd-line IT support to 120+ users, achieving ~98% uptime and secure system access.",
        "Logged and tracked incidents in Freshdesk/Zoho Desk using ITIL workflows, meeting SLAs and improving response time ~25%.",
        "Developed SQL scripts and automated validation checks that enhanced surveillance data accuracy.",
        "Administered Active Directory (provisioning, access rights, group policy) and Microsoft 365 (Outlook, Teams, SharePoint).",
        "Configured and monitored Cisco Meraki network devices for reliable, secure connectivity.",
        "Authored knowledge-base articles and delivered user training and onboarding.",
    ]},
    {"org":"PATH","title":"ICT Intern","dates":"Apr 2023 – Dec 2023","loc":"Homa Bay, Kenya","bullets":[
        "Assisted facility data-system deployment and digitalization for surveillance and reporting readiness.",
        "Supported LAN/WAN, routers, switches, and structured cabling; logged/tracked Freshdesk tickets under SLA monitoring.",
        "Deployed and configured printers, VoIP phones, and desktop imaging (PXE boot); contributed to process documentation.",
    ]},
],
"certs": CERTS_HEALTH,
"achievements": ACH,
}

# 2) VARIANT — INGO / UN ICT Officer / IT Support
ingo = {
"headline": "ICT Officer  |  Health Systems & IT Support  |  Microsoft 365 · Active Directory · Networking · ITIL/SLA",
"summary": ("ICT professional with ~3 years at PATH delivering reliable IT support and systems administration for "
    "120+ users across multi-site health programs. Strong in Microsoft 365, Active Directory, networking, and "
    "ITIL/SLA-driven service desk operations, with hands-on experience supporting mission-critical health "
    "information systems (KenyaEMR, DHIS2). Known for structured troubleshooting, documentation, and user training "
    "in fast-paced, multicultural environments. Seeking an ICT Officer / IT support role in a mission-driven organization."),
"skills": [
    ("IT Support & Service Desk", "1st/2nd-line support, Freshdesk, Zoho Desk, Jira Service Management, Zendesk, ITIL, SLA management, incident & problem management"),
    ("Systems & Identity", "Windows 10/11 & Server 2016-2022, Active Directory/Azure AD, Group Policy, Microsoft 365, Linux (Ubuntu/CentOS)"),
    ("Networking", "LAN/WAN, VLANs, VPN, DNS/DHCP, TCP/IP, Cisco Meraki, routers/switches, Wi-Fi, performance monitoring"),
    ("Security & Continuity", "Endpoint protection/MFA, IAM, patch management, backups (Veeam/rsync), disaster recovery, security-awareness training"),
    ("Health & Data", "KenyaEMR/TaifaCare, DHIS2, data quality, SQL, Power BI, reporting; IT asset management; documentation"),
],
"experience": [
    {"org":"PATH","title":"ICT / HMIS Officer","dates":"Mar 2025 – Present","loc":"Kisumu/Nyamira, Kenya (Hybrid)","bullets":[
        "Ensure availability of KenyaEMR/TaifaCare and DHIS2 for clinical users, sustaining >98% uptime across facilities.",
        "Maintain Windows/Linux servers with backups, patching, encryption, and disaster-recovery controls.",
        "Support EMR integrations and produce uptime, incident, and IT-security posture reports for decision-makers.",
        "Train facility staff and ICT officers; author SOPs, job aids, and technical documentation.",
    ]},
    {"org":"PATH","title":"ICT Associate","dates":"Jan 2024 – Mar 2025","loc":"Kisumu, Kenya","bullets":[
        "Delivered 1st/2nd-line support to 120+ users across locations, achieving ~98% uptime and ~25% faster response times.",
        "Managed incidents in Freshdesk/Zoho Desk under ITIL workflows with SLA compliance and clear escalation.",
        "Administered Active Directory and Microsoft 365 (Outlook, Teams, SharePoint); enforced least-privilege access.",
        "Configured and monitored Cisco Meraki network devices; troubleshot LAN/WAN, VPN, and connectivity issues.",
        "Supported endpoint protection, patching, and MFA; maintained IT asset inventory and knowledge base.",
    ]},
    {"org":"PATH","title":"ICT Intern","dates":"Apr 2023 – Dec 2023","loc":"Homa Bay, Kenya","bullets":[
        "Supported LAN/WAN infrastructure, desktop deployment, printers, and VoIP; assisted server-room diagnostics.",
        "Logged and tracked tickets in Freshdesk under ITIL/SLA workflows; contributed to IT documentation and audits.",
    ]},
],
"certs": CERTS_HEALTH,
"achievements": ACH,
}

# 3) VARIANT — Enterprise / Remote IT Operations & Sysadmin
itops = {
"headline": "IT Operations & Systems Administrator  |  Microsoft 365 · Active Directory · Networking · Automation  |  Remote-ready",
"summary": ("IT operations professional with ~3 years at PATH supporting 120+ users in hybrid Windows/Microsoft 365 "
    "environments. Skilled across service desk (Freshdesk/Jira, ITIL/SLA), Active Directory, networking, endpoint "
    "security, and Windows/Linux server administration, with growing automation (PowerShell, Bash, Python) and "
    "cloud (Azure) capability. Reliable, documentation-driven, and effective in remote, SLA-based support teams."),
"skills": [
    ("Service Desk & ITSM", "Freshdesk, Jira Service Management, Zendesk, Zoho Desk, ITIL, SLA, incident/RCA, knowledge base"),
    ("Systems Administration", "Windows 10/11 & Server 2016-2022, Active Directory/Azure AD, Group Policy, Microsoft 365, Linux (Ubuntu/CentOS), Apache Tomcat"),
    ("Networking", "LAN/WAN, VLANs, VPN, DNS/DHCP, TCP/IP, Cisco Meraki, routing/switching, Wi-Fi"),
    ("Cloud & Automation", "Microsoft Azure (AZ-900), Azure AD, Git/GitHub; PowerShell, Bash, Python scripting; basic Docker/AWS"),
    ("Security & Continuity", "Endpoint protection/EDR, MFA, IAM, patch management, backups (Veeam/rsync), disaster recovery"),
],
"experience": [
    {"org":"PATH","title":"ICT / Systems Officer (HMIS)","dates":"Mar 2025 – Present","loc":"Kenya (Hybrid)","bullets":[
        "Administer Windows/Linux servers hosting mission-critical applications, sustaining >98% uptime with backups, patching, and DR.",
        "Automate validation and operational reporting with SQL and Python, reducing manual effort ~40%.",
        "Monitor systems, resolve incidents, and produce uptime and IT-security posture reports.",
        "Maintain documentation, SOPs, and access controls across enterprise platforms.",
    ]},
    {"org":"PATH","title":"ICT Associate","dates":"Jan 2024 – Mar 2025","loc":"Kisumu, Kenya","bullets":[
        "Provided 1st/2nd-line support to 120+ users; handled 20+ tickets/day with ~98% uptime and ~25% faster response.",
        "Administered Active Directory and Microsoft 365; enforced least-privilege access and group policy.",
        "Configured/monitored Cisco Meraki networks; troubleshot LAN/WAN, VPN, DNS/DHCP issues.",
        "Supported endpoint protection, patching, MFA; automated routine tasks with PowerShell/Bash.",
        "Logged incidents in Freshdesk/Jira under ITIL/SLA; maintained knowledge base and asset inventory.",
    ]},
    {"org":"PATH","title":"ICT Intern","dates":"Apr 2023 – Dec 2023","loc":"Homa Bay, Kenya","bullets":[
        "Supported desktop deployment, imaging (PXE), printers, VoIP, and LAN/WAN infrastructure.",
        "Shadowed senior engineers on firewall rules, VLAN segmentation, and Linux server provisioning.",
    ]},
],
"certs": CERTS_HEALTH,
"achievements": ACH,
}

def _export_pdf(data, path):
    """PDF twin of each DOCX track (reportlab)."""
    from pdf_builder import build_cv_pdf
    from resume_theme import CONTACT as THEME_CONTACT
    build_cv_pdf(data, path, THEME_CONTACT)


if __name__ == "__main__":
    here = os.path.dirname(os.path.abspath(__file__))
    root = os.path.dirname(here)
    downloads = os.path.join(root, "downloads")
    os.makedirs(downloads, exist_ok=True)

    tracks = [
        ("Meshack_Ariri_Resume_HealthInformatics", health),
        ("Meshack_Ariri_Resume_ICT_Officer_INGO", ingo),
        ("Meshack_Ariri_Resume_IT_Operations", itops),
    ]
    for name, data in tracks:
        build(data, os.path.join(here, f"{name}.docx"))
        build(data, os.path.join(downloads, f"{name}.docx"))
        _export_pdf(data, os.path.join(here, f"{name}.pdf"))
        _export_pdf(data, os.path.join(downloads, f"{name}.pdf"))
    print("All resumes generated (DOCX + PDF) in resumes/ and downloads/.")
